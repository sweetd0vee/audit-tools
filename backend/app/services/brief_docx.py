from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.citations import CITE_RE

_BOOKMARK_IDS = 0


def _next_bookmark_id() -> str:
    global _BOOKMARK_IDS
    _BOOKMARK_IDS += 1
    return str(_BOOKMARK_IDS)


def add_hyperlink(paragraph, text: str, url: str, *, italic: bool = False) -> None:
    """Clickable URL. python-docx has no public helper for this."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    if italic:
        r_pr.append(OxmlElement("w:i"))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    r_pr.append(sz)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    r_pr.append(fonts)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_anchor_hyperlink(paragraph, text: str, anchor: str, *, size: int = 12) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    r_pr.append(sz)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(fonts)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    bid = _next_bookmark_id()
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    paragraph._p.append(start)
    paragraph._p.append(end)


def _set_run_font(run, *, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    try:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    except Exception:
        pass


def _style_paragraph(paragraph, *, first_line: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(6)
    fmt.space_before = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        fmt.first_line_indent = Cm(1.25)


def _add_plain_run(
    paragraph, text: str, *, italic: bool = False, bold: bool = False, size: int = 12
) -> None:
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, italic=italic, bold=bold)


def _set_table_borders(
    table, *, val: str = "single", sz: str = "4", color: str = "000000"
) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        if val != "nil":
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def _set_cell_borders(
    cell, *, val: str = "single", sz: str = "4", color: str = "000000"
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def _dxa(cm: float) -> str:
    return str(int(round(float(Cm(cm).twips))))


def _set_tbl_width(table, cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), _dxa(cm))
    tbl_w.set(qn("w:type"), "dxa")


def _set_tbl_grid(table, widths_cm: list[float]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), _dxa(width))
        grid.append(col)
    _set_tbl_width(table, sum(widths_cm))
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            _set_cell_width(cell, width)


def _add_program_table(doc: Document, rows: int, widths_cm: list[float], *, bordered: bool):
    table = doc.add_table(rows=rows, cols=len(widths_cm))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_tbl_grid(table, widths_cm)
    _set_table_borders(table, val="single" if bordered else "nil")
    return table


def _set_cell_width(cell, cm: float) -> None:
    cell.width = Cm(cm)
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), _dxa(cm))
    tc_w.set(qn("w:type"), "dxa")


def _fill_cell(
    cell,
    text: str,
    sources_by_n: dict[int, dict[str, Any]] | None = None,
    *,
    bold: bool = False,
    italic: bool = False,
    size: int = 12,
    center: bool = False,
    justify: bool = False,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(2)
    fmt.line_spacing = 1.15
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if sources_by_n and CITE_RE.search(text or ""):
        add_text_with_cites(paragraph, text or "", sources_by_n, size=size)
        return
    run = paragraph.add_run(text or "")
    _set_run_font(run, size=size, bold=bold, italic=italic)


def add_text_with_cites(
    paragraph, text: str, sources_by_n: dict[int, dict[str, Any]], *, size: int = 12
) -> None:
    """Render a paragraph, turning [n] into links to the appendix bookmark."""
    pos = 0
    for match in CITE_RE.finditer(text or ""):
        if match.start() > pos:
            _add_plain_run(paragraph, text[pos : match.start()], size=size)
        n = int(match.group(1))
        if n in sources_by_n:
            add_anchor_hyperlink(paragraph, match.group(0), f"cite_{n}", size=size)
        else:
            _add_plain_run(paragraph, match.group(0), size=size)
        pos = match.end()
    if pos < len(text or ""):
        _add_plain_run(paragraph, text[pos:], size=size)


def _strip_md(line: str) -> str:
    line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
    line = re.sub(r"`([^`]+)`", r"\1", line)
    return line.strip()


def add_markdown_block(doc: Document, md: str, sources_by_n: dict[int, dict[str, Any]]) -> None:
    for raw in (md or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            p = doc.add_heading(_strip_md(line[4:]), level=3)
            for run in p.runs:
                _set_run_font(run, size=13, bold=True)
            continue
        if line.startswith("## "):
            p = doc.add_heading(_strip_md(line[3:]), level=2)
            for run in p.runs:
                _set_run_font(run, size=14, bold=True)
            continue
        if line.startswith("# "):
            p = doc.add_heading(_strip_md(line[2:]), level=1)
            for run in p.runs:
                _set_run_font(run, size=16, bold=True)
            continue
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _style_paragraph(p)
            add_text_with_cites(p, _strip_md(line[2:]), sources_by_n)
            continue
        numbered = re.match(r"^(\d+)[.)]\s+(.*)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            _style_paragraph(p)
            add_text_with_cites(p, _strip_md(numbered.group(2)), sources_by_n)
            continue
        p = doc.add_paragraph()
        _style_paragraph(p, first_line=True)
        add_text_with_cites(p, _strip_md(line), sources_by_n)


def _new_doc_with_title(
    *,
    footer_text: str,
    title_text: str,
    inspection_name: str,
    period: str | None,
    keywords: list[str],
    note_text: str,
) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(footer_text)
    _set_run_font(fr, size=9, italic=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(title_text)
    _set_run_font(tr, size=20, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(inspection_name)
    _set_run_font(sr, size=14, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_s = period or "не указан"
    kws = ", ".join(keywords) if keywords else "—"
    mr = meta.add_run(f"Период: {period_s}. Ключевые слова: {kws}")
    _set_run_font(mr, size=11, italic=True)

    note = doc.add_paragraph()
    _style_paragraph(note)
    nr = note.add_run(note_text)
    _set_run_font(nr, size=11, italic=True)
    return doc


def _add_sources_section(
    doc: Document,
    sources: list[dict[str, Any]],
    *,
    title: str,
    intro_text: str,
    include_filename_fallback: bool = True,
) -> None:
    if not sources:
        return
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        _set_run_font(run, size=16, bold=True)
    intro = doc.add_paragraph()
    _style_paragraph(intro)
    ir = intro.add_run(intro_text)
    _set_run_font(ir, size=11, italic=True)
    for src in sources:
        n = int(src["n"])
        p = doc.add_paragraph()
        _style_paragraph(p)
        add_bookmark(p, f"cite_{n}")
        _add_plain_run(p, f"[{n}] ", bold=True)
        article = (src.get("article") or "").strip()
        src_title = (src.get("title") or "акт").strip()
        if article:
            _add_plain_run(p, f"{src_title} — {article}. ")
        else:
            _add_plain_run(p, f"{src_title}. ")
        url = (src.get("url") or "").strip()
        if url:
            add_hyperlink(p, url, url)
        elif include_filename_fallback and src.get("filename"):
            _add_plain_run(p, f"файл: {src['filename']}", italic=True)
        excerpt = (src.get("excerpt") or "").strip()
        if excerpt:
            q = doc.add_paragraph()
            _style_paragraph(q)
            q.paragraph_format.left_indent = Cm(1.0)
            _add_plain_run(q, excerpt, italic=True)


_PROGRAM_NOTE = (
    "При необходимости вопросы, подлежащие аудиту, могут быть изменены и уточнены "
    "руководителем проверки по согласованию с директором Департамента внутреннего аудита."
)

_PROGRAM_FONT = 14
_PROGRAM_LABEL_W = 5.10
_PROGRAM_VALUE_W = 11.90
_PROGRAM_NUM_W = 1.50
_PROGRAM_Q_W = 15.50
_PROGRAM_SIGN_L = 7.63
_PROGRAM_SIGN_R = 3.08


def _set_program_document_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(_PROGRAM_FONT)


def write_program_docx(
    path: Path,
    *,
    inspection_name: str,
    period: str | None,
    keywords: list[str],
    case_id: str,
    body: str,
    sources: list[dict[str, Any]],
    questions: list[str] | None = None,
) -> Path:
    global _BOOKMARK_IDS
    _BOOKMARK_IDS = 0
    _ = body
    _ = keywords
    _ = case_id

    doc = Document()
    _set_program_document_font(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    tr = title.add_run("ПРОГРАММА")
    _set_run_font(tr, size=_PROGRAM_FONT, bold=True)

    sources_by_n = {int(s["n"]): s for s in sources}
    info_widths = [_PROGRAM_LABEL_W, _PROGRAM_VALUE_W]
    q_widths = [_PROGRAM_NUM_W, _PROGRAM_Q_W]

    info = _add_program_table(doc, 5, info_widths, bordered=True)
    info_rows = [
        ("Название проверки", inspection_name or ""),
        ("Аудируемый период", period or "уточняется"),
        ("Сроки проведения", ""),
        ("Руководитель проверки", ""),
        ("Члены рабочей группы", ""),
    ]
    for row, (label, value) in zip(info.rows, info_rows):
        _fill_cell(row.cells[0], label, size=_PROGRAM_FONT)
        _fill_cell(row.cells[1], value, sources_by_n, size=_PROGRAM_FONT)

    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(6)
    gap.paragraph_format.space_after = Pt(6)

    q_rows = questions or [""]
    qtable = _add_program_table(doc, 1 + len(q_rows), q_widths, bordered=True)
    _fill_cell(qtable.rows[0].cells[0], "№ п/п", size=_PROGRAM_FONT, center=True)
    _fill_cell(
        qtable.rows[0].cells[1],
        "Вопросы, подлежащие аудиту",
        size=_PROGRAM_FONT,
        center=True,
    )
    for idx, question in enumerate(q_rows, start=1):
        row = qtable.rows[idx]
        _fill_cell(row.cells[0], f"{idx}.", size=_PROGRAM_FONT, center=True)
        _fill_cell(
            row.cells[1],
            question,
            sources_by_n,
            size=_PROGRAM_FONT,
            justify=True,
        )

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    note.paragraph_format.first_line_indent = Cm(1.0)
    note.paragraph_format.space_before = Pt(10)
    nr = note.add_run(_PROGRAM_NOTE)
    _set_run_font(nr, size=_PROGRAM_FONT)

    spacer = doc.add_paragraph()
    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sign = _add_program_table(
        doc, 1, [_PROGRAM_SIGN_L, _PROGRAM_SIGN_R], bordered=False
    )
    _fill_cell(
        sign.rows[0].cells[0],
        "Менеджер по направлению деятельности",
        size=_PROGRAM_FONT,
    )
    _fill_cell(sign.rows[0].cells[1], "", size=_PROGRAM_FONT)

    _add_sources_section(
        doc,
        sources,
        title="Источники: статьи и фрагменты",
        intro_text=(
            "Каждая ссылка [n] в тексте указывает на фрагмент ниже. "
            "Официальный URL — страница, с которой акт скачан в библиотеку кейса."
        ),
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def write_total_docx(
    path: Path,
    *,
    inspection_name: str,
    period: str | None,
    keywords: list[str],
    case_id: str,
    body: str,
    sources: list[dict[str, Any]],
) -> Path:
    global _BOOKMARK_IDS
    _BOOKMARK_IDS = 0

    doc = _new_doc_with_title(
        footer_text=(
            f"Саммари total · {inspection_name} · кейс {case_id} · знания модели · черновик"
        ),
        title_text="Конспект по теме (знания модели)",
        inspection_name=inspection_name,
        period=period,
        keywords=keywords,
        note_text=(
            "Краткий конспект самого важного по теме из знаний языковой модели, "
            "без опоры на скачанные акты базы знаний кейса. "
            "Номера [n] ведут к списку актов и статей в конце. "
            "Редакции и точные формулировки нужно сверять с первоисточником. "
            "Это черновик, не аудиторское суждение."
        ),
    )

    sources_by_n = {int(s["n"]): s for s in sources}
    add_markdown_block(doc, body, sources_by_n)

    _add_sources_section(
        doc,
        sources,
        title="Источники: акты и статьи",
        intro_text=(
            "Ссылки [n] в тексте указывают на акт ниже. "
            "URL — если модель его помнит; иначе сверяйте на pravo.by / nbrb.by."
        ),
        include_filename_fallback=False,
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def write_brief_docx(
    path: Path,
    *,
    inspection_name: str,
    period: str | None,
    keywords: list[str],
    case_id: str,
    overview: str,
    chapters: list[dict[str, Any]],
    sources: list[dict[str, Any]],
) -> Path:
    global _BOOKMARK_IDS
    _BOOKMARK_IDS = 0

    doc = _new_doc_with_title(
        footer_text=f"Саммари НПА · {inspection_name} · кейс {case_id} · черновик",
        title_text="Саммари нормативной базы",
        inspection_name=inspection_name,
        period=period,
        keywords=keywords,
        note_text=(
            "Карточка существенного по каждому акту: только нормы, которые важны "
            "для этой проверки, не перечень всех статей. "
            "Оглавление и нормы — из текста документов. "
            "Это черновик, не аудиторское суждение."
        ),
    )

    if overview.strip():
        h = doc.add_heading("Обзор проверки", level=1)
        for run in h.runs:
            _set_run_font(run, size=16, bold=True)
        add_markdown_block(doc, overview, {int(s["n"]): s for s in sources})

    sources_by_n = {int(s["n"]): s for s in sources}
    for chapter in chapters:
        h = doc.add_heading(chapter.get("title") or "Акт", level=1)
        for run in h.runs:
            _set_run_font(run, size=16, bold=True)
        add_markdown_block(doc, chapter.get("body") or "", sources_by_n)

    _add_sources_section(
        doc,
        sources,
        title="Источники: статьи и фрагменты",
        intro_text=(
            "Каждая ссылка [n] в тексте указывает на фрагмент ниже. "
            "Официальный URL — страница, с которой акт скачан в библиотеку кейса."
        ),
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
