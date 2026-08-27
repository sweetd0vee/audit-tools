from __future__ import annotations

import math
import re
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.services.brief_docx import (
    _add_opinion_paragraph,
    _add_program_table,
    _set_cell_borders,
    _set_cell_width,
    _set_document_base_font,
    _set_run_font,
    _set_table_borders,
    _set_tbl_grid,
    _strip_md,
    add_bookmark,
    add_opinion_markdown,
)

_FONT_SIZE = 14
_TITLE = "Аудиторское заключение (черновик)"
_COVER_TITLE = "Аудиторское заключение"
_COVER_DEPT = "Департамент внутреннего аудита"
_COVER_IMAGE = Path(__file__).resolve().parent.parent / "assets" / "conclusion_cover.png"
_COVER_DRAWING_ID = 100
_TOC_HEADING = "Содержание"
_SECTION_I = "Аудиторское мнение по итогам проверки."
_SECTION_II = "Основные результаты аудита и итоговые аудиторские рекомендации."
_SECTION_III = "Наблюдения по итогам проверки."
_SECTION_III_TAX = (
    "Оценка соответствия деятельности принципам налогообложения и защиты прав плательщиков."
)
_SECTION_LAST = "Общая информация об аудиторской проверке."
_ROMAN = [
    "",
    "I",
    "II",
    "III",
    "IV",
    "V",
    "VI",
    "VII",
    "VIII",
    "IX",
    "X",
    "XI",
    "XII",
    "XIII",
    "XIV",
    "XV",
]
_MATERIALITY = {"высокий", "средний", "низкий"}
_HEADING_PREFIX = re.compile(r"^#{1,6}\s*")
_SECTION_RE = re.compile(
    r"^(?:Раздел\s+)?([IVX]{1,6})\.?\s*(.*)$",
    re.I,
)
_OBS_RE = re.compile(
    r"^Наблюдение\s+(\d+(?:\.\d+)?)\.?\s*(.*)$",
    re.I,
)
_MATERIALITY_RE = re.compile(
    r"^(?:уровень\s+)?существенност[иь]:\s*(высокий|средний|низкий)\b",
    re.I,
)
_HYP_RE = re.compile(r"^гипотеза:\s*(\d+)\b", re.I)
_REC_RE = re.compile(r"^(?:аудиторская\s+)?рекомендация:\s*(.*)$", re.I)
_SKIP_FIELD_RE = re.compile(
    r"^(аудитор|объект аудита|руководитель объекта|срок|программа)\b",
    re.I,
)
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")
_TAX_HINTS = ("налог", "ндс", "плательщик", "налогооблож")
_GENERAL_LABELS = (
    "Основание проведения аудита",
    "Срок проведения",
    "Группа аудиторов",
    "Вид аудита",
    "Дата составления заключения",
)
_DROPPED_GENERAL_LABELS = ("Аудируемый период",)


@dataclass
class Observation:
    number: str
    title: str
    body: str
    materiality: str
    recommendation: str
    hypothesis_n: str = ""


@dataclass
class ReportSection:
    roman: str
    title: str
    intro: str = ""
    observations: list[Observation] = field(default_factory=list)
    general_items: list[tuple[str, str]] = field(default_factory=list)
    kind: str = "observations"


@dataclass
class ConclusionDocument:
    sections: list[ReportSection] = field(default_factory=list)


def roman_numeral(n: int) -> str:
    if 1 <= n < len(_ROMAN):
        return _ROMAN[n]
    return str(n)


def normalize_materiality(value: str | None, fallback: str = "средний") -> str:
    raw = (value or "").strip().lower()
    if raw in _MATERIALITY:
        return raw
    if "высок" in raw:
        return "высокий"
    if "низк" in raw:
        return "низкий"
    if "средн" in raw:
        return "средний"
    fb = (fallback or "").strip().lower()
    if fb in _MATERIALITY:
        return fb
    if "высок" in fb:
        return "высокий"
    if "низк" in fb:
        return "низкий"
    return "средний"


def materiality_from_priority(priority: str | None) -> str:
    return normalize_materiality(priority, "средний")


def _bare(line: str) -> str:
    return _HEADING_PREFIX.sub("", (line or "").strip())


def _is_toc_heading(text: str) -> bool:
    low = text.lower()
    return "содержан" in low or "разделы аудиторского заключения" in low


def _is_section_i(title: str) -> bool:
    return "аудиторское мнение" in (title or "").lower()


def _is_section_ii(title: str) -> bool:
    return "основные результаты" in (title or "").lower()


def _is_tax_template_title(title: str) -> bool:
    low = (title or "").lower()
    return "принципам налогообложения" in low or "защиты прав плательщиков" in low


def _inspection_is_tax(inspection: str | None) -> bool:
    low = (inspection or "").lower()
    return any(hint in low for hint in _TAX_HINTS)


def default_section_iii_title(inspection: str | None = None) -> str:
    if _inspection_is_tax(inspection):
        return _SECTION_III_TAX
    return _SECTION_III


def _dot_title(title: str) -> str:
    heading = (title or "").strip()
    if heading and not heading.endswith("."):
        heading += "."
    return heading


def _is_general(title: str) -> bool:
    return "общая информация" in (title or "").lower()


def _clean_paragraphs(text: str) -> str:
    lines = [ln.rstrip() for ln in (text or "").splitlines()]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return "\n".join(lines).strip()


def default_general_items() -> list[tuple[str, str]]:
    return [
        (
            "Основание проведения аудита",
            "уточняется (приказ / план работы службы внутреннего аудита — заполняет аудитор)",
        ),
        ("Срок проведения", "уточняется"),
        ("Группа аудиторов", "заполняет аудитор"),
        ("Вид аудита", "Тематическая аудиторская проверка"),
        ("Дата составления заключения", "заполняет аудитор"),
    ]


def _parse_observation_block(
    number: str,
    title: str,
    lines: list[str],
    *,
    fallback_materiality: str = "средний",
) -> Observation:
    materiality = fallback_materiality
    hypothesis_n = ""
    rec_lines: list[str] = []
    body_lines: list[str] = []
    in_rec = False
    for raw in lines:
        stripped = raw.strip()
        if not stripped:
            if in_rec:
                rec_lines.append("")
            elif body_lines:
                body_lines.append("")
            continue
        bare = _bare(stripped)
        match_m = _MATERIALITY_RE.match(bare)
        if match_m:
            materiality = normalize_materiality(match_m.group(1), fallback_materiality)
            continue
        match_h = _HYP_RE.match(bare)
        if match_h:
            hypothesis_n = match_h.group(1)
            continue
        match_r = _REC_RE.match(bare)
        if match_r:
            in_rec = True
            rest = (match_r.group(1) or "").strip()
            if rest:
                rec_lines.append(rest)
            continue
        if _SKIP_FIELD_RE.match(bare):
            continue
        if in_rec:
            rec_lines.append(stripped)
        else:
            body_lines.append(raw)
    title_clean = _strip_md(title).strip(" .")
    return Observation(
        number=number,
        title=title_clean,
        body=_clean_paragraphs("\n".join(body_lines)),
        materiality=normalize_materiality(materiality, fallback_materiality),
        recommendation=_clean_paragraphs("\n".join(rec_lines)),
        hypothesis_n=hypothesis_n,
    )


def _parse_general_items(lines: list[str]) -> list[tuple[str, str]]:
    items: list[tuple[str, str]] = []
    current_label = ""
    current_lines: list[str] = []
    known_labels = _GENERAL_LABELS + _DROPPED_GENERAL_LABELS

    def flush() -> None:
        nonlocal current_label, current_lines
        if current_label:
            items.append(
                (current_label, _clean_paragraphs("\n".join(current_lines)) or "уточняется")
            )
        current_label = ""
        current_lines = []

    for raw in lines:
        stripped = raw.strip()
        bare = _bare(stripped)
        label = next((lab for lab in known_labels if bare.lower().startswith(lab.lower())), None)
        if label:
            flush()
            if label in _DROPPED_GENERAL_LABELS:
                continue
            current_label = label
            rest = bare[len(label) :].lstrip(" .:–-")
            current_lines = [rest] if rest else []
            continue
        if current_label:
            current_lines.append(raw)
    flush()
    if not items:
        return default_general_items()
    have = {label.lower() for label, _ in items}
    for label, value in default_general_items():
        if label.lower() not in have:
            items.append((label, value))
    return items


def fallback_from_hypotheses(
    hypotheses: list[dict[str, str]],
    *,
    leftover: str = "",
    inspection_name: str | None = None,
) -> ConclusionDocument:
    observations: list[Observation] = []
    leftover_body = _clean_paragraphs(leftover)
    for i, row in enumerate(hypotheses, start=1):
        n = str(row.get("n") or i)
        title = (row.get("hypothesis") or f"Гипотеза {n}").strip()
        if len(title) > 140:
            title = title[:137].rstrip() + "…"
        parts: list[str] = []
        if leftover_body and i == 1:
            parts.append(leftover_body)
        else:
            hyp = (row.get("hypothesis") or "").strip()
            if hyp:
                parts.append(
                    "По итогам рассмотрения подтверждённой гипотезы отмечается следующий "
                    "риск / недостаток контроля. " + hyp.rstrip(".") + "."
                )
            risk = (row.get("risk") or "").strip()
            why = (row.get("why_risk") or "").strip()
            if risk or why:
                parts.append(" ".join(p for p in (risk, why) if p))
            how = (row.get("how_to_test") or "").strip()
            npa = (row.get("npa_criteria") or "").strip()
            extra: list[str] = []
            if how:
                extra.append(f"В ходе проверки целесообразно опереться на процедуру: {how.rstrip('.')}.")
            if npa:
                extra.append(f"Критерии: {npa.rstrip('.')}.")
            if extra:
                parts.append(" ".join(extra))
        rec = (row.get("working_paper") or "").strip()
        if not rec:
            rec = (
                "Проработать выявленный риск по подтверждённой гипотезе и закрепить "
                "контрольные процедуры в локальных актах и операционном контуре банка."
            )
        observations.append(
            Observation(
                number=f"3.{i}",
                title=title.rstrip("."),
                body=_clean_paragraphs("\n\n".join(parts)),
                materiality=materiality_from_priority(row.get("priority")),
                recommendation=rec,
                hypothesis_n=n,
            )
        )
    return ConclusionDocument(
        sections=[
            ReportSection(
                roman="III",
                title=default_section_iii_title(inspection_name).rstrip("."),
                observations=observations,
                kind="observations",
            ),
            ReportSection(
                roman="IV",
                title=_SECTION_LAST.rstrip("."),
                general_items=default_general_items(),
                kind="general",
            ),
        ]
    )


def parse_conclusion_markdown(
    md: str,
    *,
    hypotheses: list[dict[str, str]] | None = None,
    inspection_name: str | None = None,
) -> ConclusionDocument:
    hypotheses = hypotheses or []
    lines = (md or "").splitlines()
    blocks: list[tuple[str, str, str, list[str]]] = []
    kind, code, title, buf = "lead", "", "", []

    def push() -> None:
        if kind != "lead" or any(ln.strip() for ln in buf):
            blocks.append((kind, code, title, list(buf)))

    for raw in lines:
        stripped = raw.strip()
        bare = _bare(stripped)
        if _is_toc_heading(bare):
            push()
            kind, code, title, buf = "toc", "", bare, []
            continue
        obs = _OBS_RE.match(bare)
        if obs:
            push()
            kind, code, title, buf = "obs", obs.group(1), (obs.group(2) or "").strip(), []
            continue
        sec = _SECTION_RE.match(bare)
        looks_heading = (
            stripped.startswith("#")
            or bare.lower().startswith("раздел ")
            or (sec and sec.group(1).isupper() and (stripped.startswith("#") or len(bare) < 120))
        )
        if sec and looks_heading:
            roman = sec.group(1).upper()
            rest = (sec.group(2) or "").strip()
            if roman in {"I", "II"} or _is_section_i(rest) or _is_section_ii(rest):
                push()
                kind, code, title, buf = "skip", roman, rest, []
                continue
            if kind == "toc" and not stripped.startswith("#") and not bare.lower().startswith("раздел "):
                buf.append(raw)
                continue
            push()
            kind, code, title, buf = "sec", roman, rest, []
            continue
        buf.append(raw)
    push()

    sections: list[ReportSection] = []
    current: ReportSection | None = None
    leftovers: list[str] = []
    used_hyps = 0

    def close_current() -> None:
        nonlocal current
        if current is None:
            return
        if current.kind == "observations" and not current.observations and current.intro:
            current.observations.append(
                Observation(
                    number=f"{_roman_to_int(current.roman)}.1",
                    title=current.title,
                    body=current.intro,
                    materiality="средний",
                    recommendation="",
                )
            )
            current.intro = ""
        sections.append(current)
        current = None

    for kind, code, title, body_lines in blocks:
        if kind in {"lead", "toc", "skip"}:
            leftovers.extend(body_lines)
            continue
        if kind == "sec":
            close_current()
            if _is_general(title):
                current = ReportSection(
                    roman=code or "IV",
                    title=title or _SECTION_LAST.rstrip("."),
                    general_items=_parse_general_items(body_lines),
                    kind="general",
                )
                close_current()
                continue
            current = ReportSection(
                roman=code or roman_numeral(len(sections) + 3),
                title=(title or default_section_iii_title(inspection_name)).rstrip("."),
                intro=_clean_paragraphs("\n".join(body_lines)),
                kind="observations",
            )
            continue
        if kind == "obs":
            if current is None or current.kind != "observations":
                close_current()
                roman = code.split(".", 1)[0] if "." in code else roman_numeral(len(sections) + 3)
                current = ReportSection(
                    roman=roman_numeral(int(roman)) if roman.isdigit() else roman,
                    title=_SECTION_III.rstrip("."),
                    kind="observations",
                )
            fallback = "средний"
            if used_hyps < len(hypotheses):
                fallback = materiality_from_priority(hypotheses[used_hyps].get("priority"))
                used_hyps += 1
            current.observations.append(
                _parse_observation_block(code, title, body_lines, fallback_materiality=fallback)
            )
    close_current()

    observation_sections = [s for s in sections if s.kind == "observations" and s.observations]
    if not observation_sections:
        leftover_text = _clean_paragraphs("\n".join(leftovers))
        return fallback_from_hypotheses(
            hypotheses, leftover=leftover_text, inspection_name=inspection_name
        )

    return ConclusionDocument(
        sections=_canonicalize_report(
            sections, inspection_name=inspection_name
        )
    )


def _roman_to_int(value: str) -> int:
    raw = (value or "").upper().strip()
    if raw.isdigit():
        return int(raw)
    try:
        return _ROMAN.index(raw)
    except ValueError:
        return 3


def _canonicalize_report(
    sections: list[ReportSection],
    *,
    inspection_name: str | None = None,
) -> list[ReportSection]:
    observations: list[Observation] = []
    intro_parts: list[str] = []
    general: ReportSection | None = None
    first_title = ""
    for section in sections:
        if section.kind == "general":
            general = section
            continue
        if not first_title and section.title:
            first_title = section.title
        if section.intro:
            intro_parts.append(section.intro)
        observations.extend(section.observations)
    if general is None:
        general = ReportSection(
            roman="IV",
            title=_SECTION_LAST.rstrip("."),
            general_items=default_general_items(),
            kind="general",
        )
    else:
        general.roman = "IV"
        general.title = _SECTION_LAST.rstrip(".")
    title = first_title.strip() or default_section_iii_title(inspection_name)
    if _is_tax_template_title(title) and not _inspection_is_tax(inspection_name):
        title = default_section_iii_title(inspection_name)
    obs = ReportSection(
        roman="III",
        title=title.rstrip("."),
        intro=_clean_paragraphs("\n\n".join(intro_parts)),
        observations=observations,
        kind="observations",
    )
    for i, item in enumerate(obs.observations, start=1):
        item.number = f"3.{i}"
    return [obs, general]


def iter_observations(report: ConclusionDocument) -> list[Observation]:
    out: list[Observation] = []
    for section in report.sections:
        out.extend(section.observations)
    return out


def missing_hypothesis_rows(
    report: ConclusionDocument,
    hypotheses: list[dict[str, str]],
) -> list[dict[str, str]]:
    unused = list(hypotheses)
    unused_ns = [str(row.get("n") or "") for row in unused]
    for obs in iter_observations(report):
        n = str(obs.hypothesis_n or "").strip()
        if n and n in unused_ns:
            idx = unused_ns.index(n)
            unused.pop(idx)
            unused_ns.pop(idx)
        elif unused:
            row = unused.pop(0)
            unused_ns.pop(0)
            obs.hypothesis_n = str(row.get("n") or obs.hypothesis_n)
    return unused


def ensure_all_hypotheses(
    report: ConclusionDocument,
    hypotheses: list[dict[str, str]],
    *,
    inspection_name: str | None = None,
) -> ConclusionDocument:
    missing = missing_hypothesis_rows(report, hypotheses)
    if not missing:
        if iter_observations(report):
            return ConclusionDocument(
                sections=_canonicalize_report(
                    report.sections, inspection_name=inspection_name
                )
            )
        return fallback_from_hypotheses(
            hypotheses, inspection_name=inspection_name
        )
    extra = fallback_from_hypotheses(
        missing, inspection_name=inspection_name
    )
    extra_obs = iter_observations(extra)
    sections = list(report.sections)
    obs_sec = next((section for section in sections if section.kind == "observations"), None)
    if obs_sec is None:
        sections = extra.sections + [
            section for section in sections if section.kind == "general"
        ]
    else:
        obs_sec.observations.extend(extra_obs)
    return ConclusionDocument(
        sections=_canonicalize_report(
            sections, inspection_name=inspection_name
        )
    )


def toc_entries(doc: ConclusionDocument | None = None, **_: object) -> list[tuple[str, str]]:
    title = _SECTION_III
    if doc is not None:
        for section in doc.sections:
            if section.kind == "observations" and section.title:
                title = _dot_title(section.title)
                break
    return [
        ("I", _SECTION_I),
        ("II", _SECTION_II),
        ("III", title if title.endswith(".") else title + "."),
        ("IV", _SECTION_LAST),
    ]


_LINES_PER_PAGE = 22
_CHARS_PER_LINE = 62


def _plain_len(text: str) -> int:
    cleaned = _strip_md(text or "")
    cleaned = re.sub(r"\|+", " ", cleaned)
    return len(re.sub(r"\s+", " ", cleaned).strip())


def _lines_for_text(text: str) -> float:
    lines = 0.0
    for raw in (text or "").splitlines():
        stripped = raw.strip()
        if not stripped:
            lines += 0.4
            continue
        if stripped.startswith("|"):
            lines += 1.0
            continue
        width = _plain_len(stripped) or 1
        lines += max(1.0, math.ceil(width / _CHARS_PER_LINE))
    return lines


def estimate_toc_pages(
    opinion_body: str,
    report: ConclusionDocument | None,
) -> dict[str, int]:
    """Page numbers for the TOC: cover=1, contents=2, body follows."""
    observations: list[Observation] = []
    intro = ""
    if report is not None:
        for section in report.sections:
            if section.kind == "observations":
                intro = section.intro or ""
                observations = list(section.observations)
                break
    toc_lines = 8 + len(observations)
    toc_pages = 1 if toc_lines <= 28 else 2
    cursor = 2 + toc_pages

    pages: dict[str, int] = {"I": cursor}
    opinion_lines = _lines_for_text(opinion_body) + 3
    cursor += max(1, math.ceil(opinion_lines / _LINES_PER_PAGE))

    pages["III"] = cursor
    cursor += max(0.25, _lines_for_text(intro) / _LINES_PER_PAGE)
    for observation in observations:
        pages[observation.number] = int(cursor)
        chunk_lines = (
            10
            + _lines_for_text(observation.body)
            + _lines_for_text(observation.recommendation)
        )
        cursor += max(0.55, chunk_lines / _LINES_PER_PAGE)
    pages["IV"] = max(int(cursor), pages["III"])
    return pages


def _add_runs(paragraph, parts: list[tuple[str, dict]], *, font: str) -> None:
    for text, opts in parts:
        if not text:
            continue
        run = paragraph.add_run(text)
        _set_run_font(
            run,
            size=int(opts.get("size", _FONT_SIZE)),
            bold=bool(opts.get("bold")),
            italic=bool(opts.get("italic")),
            font=font,
        )
        if opts.get("underline"):
            run.underline = True


def _add_styled_paragraph(
    doc: Document,
    *,
    font: str,
    align: str = "justify",
    space_before: int = 0,
    space_after: int = 8,
    first_line: bool = False,
) -> object:
    _ = font
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = Cm(0)
    elif align == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Cm(0)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Cm(1.25) if first_line else Cm(0)
    return paragraph


def _blank(doc: Document, *, font: str, after: int = 0) -> None:
    p = _add_styled_paragraph(doc, font=font, align="left", space_after=after)
    run = p.add_run("")
    _set_run_font(run, size=_FONT_SIZE, font=font)


def _add_body_markdown(doc: Document, md: str, *, font: str) -> None:
    lines = (md or "").splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            block: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            _add_markdown_table(doc, block, font=font)
            continue
        if stripped.startswith("```"):
            i += 1
            scheme: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                if lines[i].strip():
                    scheme.append(lines[i].rstrip())
                i += 1
            if i < len(lines) and lines[i].strip().startswith("```"):
                i += 1
            if scheme:
                _add_scheme_block(doc, scheme, font=font)
            continue
        low = stripped.lower()
        if low.startswith("схема:") or low.startswith("**схема"):
            caption = _strip_md(stripped)
            _add_opinion_paragraph(
                doc, caption, font=font, size=_FONT_SIZE, italic=True, first_line=False, space_after=4
            )
            i += 1
            continue
        if stripped.startswith("### "):
            _add_opinion_paragraph(
                doc,
                _strip_md(stripped[4:]),
                font=font,
                size=_FONT_SIZE,
                bold=True,
                first_line=False,
                space_before=8,
                space_after=6,
            )
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            _add_opinion_paragraph(
                doc,
                re.sub(r"^\d+\.\s+", "", stripped),
                font=font,
                size=_FONT_SIZE,
                first_line=False,
                bullet=True,
                space_after=4,
            )
            i += 1
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            _add_opinion_paragraph(
                doc,
                stripped[2:].strip(),
                font=font,
                size=_FONT_SIZE,
                first_line=False,
                bullet=True,
                space_after=4,
            )
            i += 1
            continue
        _add_opinion_paragraph(doc, stripped, font=font, size=_FONT_SIZE)
        i += 1


def _split_table_row(line: str) -> list[str]:
    raw = line.strip().strip("|")
    return [_strip_md(cell).strip() for cell in raw.split("|")]


def _add_markdown_table(doc: Document, lines: list[str], *, font: str) -> None:
    rows = [line for line in lines if line.strip() and not _TABLE_SEP_RE.match(line)]
    parsed = [_split_table_row(line) for line in rows]
    parsed = [row for row in parsed if any(cell for cell in row)]
    if not parsed:
        return
    cols = max(len(row) for row in parsed)
    if cols < 1:
        return
    for row in parsed:
        while len(row) < cols:
            row.append("")
        del row[cols:]
    usable = 16.5
    widths = [round(usable / cols, 2)] * cols
    widths[-1] = round(usable - sum(widths[:-1]), 2)
    table = _add_program_table(doc, len(parsed), widths, bordered=True)
    for r, row in enumerate(parsed):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.space_before = Pt(2)
            run = paragraph.add_run(value)
            _set_run_font(run, size=12, bold=(r == 0), font=font)
    _blank(doc, font=font, after=6)


def _add_scheme_block(doc: Document, lines: list[str], *, font: str) -> None:
    _add_opinion_paragraph(
        doc, "Схема", font=font, size=_FONT_SIZE, italic=True, first_line=False, space_after=4
    )
    for line in lines:
        _add_opinion_paragraph(
            doc, line, font=font, size=_FONT_SIZE, first_line=False, space_after=2
        )
    _blank(doc, font=font, after=6)


def _set_row_height(row, twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:trHeight"))
    if existing is not None:
        tr_pr.remove(existing)
    el = OxmlElement("w:trHeight")
    el.set(qn("w:val"), str(twips))
    tr_pr.append(el)


def _set_tbl_center(table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _set_paragraph_borders(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)
    p_bdr = OxmlElement("w:pBdr")
    for edge, space in (("top", "1"), ("left", "4"), ("bottom", "1"), ("right", "4")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), space)
        el.set(qn("w:color"), "auto")
        p_bdr.append(el)
    p_pr.append(p_bdr)


def _add_boxed_paragraph(
    doc: Document,
    *,
    font: str,
    first_line: bool = False,
    space_before: int = 0,
):
    paragraph = _add_styled_paragraph(
        doc,
        font=font,
        align="justify",
        space_before=space_before,
        space_after=0,
        first_line=first_line,
    )
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _set_paragraph_borders(paragraph)
    return paragraph


def _add_observation_header(doc: Document, observation: Observation, *, font: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    _set_tbl_grid(table, [4.5, 13.0])
    _set_table_borders(table, val="nil")
    _set_tbl_center(table)
    row = table.rows[0]
    _set_row_height(row, 795)
    left, right = row.cells
    _set_cell_width(left, 4.5)
    _set_cell_width(right, 13.0)
    _set_cell_borders(left, val="nil")
    _set_cell_borders(right, val="nil")
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_after = Pt(0)
    lp.paragraph_format.space_before = Pt(0)
    run = lp.add_run(f"Наблюдение {observation.number}. ")
    _set_run_font(run, size=_FONT_SIZE, font=font)
    add_bookmark(lp, f"obs_{observation.number.replace('.', '_')}")

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rp.paragraph_format.space_after = Pt(0)
    rp.paragraph_format.space_before = Pt(0)
    title_run = rp.add_run(observation.title or "Название наблюдения")
    _set_run_font(title_run, size=_FONT_SIZE, italic=True, font=font)
    title_run.underline = True


def _add_observation_summary(doc: Document, observation: Observation, *, font: str) -> None:
    materiality = _add_boxed_paragraph(doc, font=font, first_line=True, space_before=6)
    _add_runs(
        materiality,
        [
            ("Уровень существенности:", {}),
            (f" {observation.materiality}", {"bold": True}),
        ],
        font=font,
    )

    auditor = _add_boxed_paragraph(doc, font=font)
    _add_runs(
        auditor,
        [("\t", {"bold": True, "italic": True}), ("Аудитор:", {})],
        font=font,
    )

    obj = _add_boxed_paragraph(doc, font=font)
    _add_runs(obj, [("          ", {"bold": True}), ("Объект аудита: ", {})], font=font)

    head = _add_boxed_paragraph(doc, font=font, first_line=True)
    _add_runs(head, [("Руководитель объекта аудита: ", {})], font=font)

    _add_boxed_paragraph(doc, font=font, first_line=True)

    rec_label = _add_boxed_paragraph(doc, font=font)
    _add_runs(rec_label, [("          Аудиторская рекомендация: ", {"bold": True})], font=font)

    rec_parts = [ln.strip() for ln in (observation.recommendation or "").splitlines() if ln.strip()]
    if not rec_parts:
        rec_parts = [""]
    for part in rec_parts:
        rec = _add_boxed_paragraph(doc, font=font, first_line=True)
        if part:
            _add_runs(rec, [(part, {})], font=font)

    deadline = _add_boxed_paragraph(doc, font=font, first_line=True)
    _add_runs(deadline, [("Срок – ", {"bold": True})], font=font)
    _blank(doc, font=font)


def _add_roman_heading(
    doc: Document, roman: str, title: str, *, font: str, bookmark: str = ""
) -> object:
    p = _add_styled_paragraph(
        doc, font=font, align="left", space_before=12, space_after=10, first_line=False
    )
    heading = _dot_title(title)
    _add_runs(
        p,
        [
            (f"{roman}.", {"bold": True, "size": 16}),
            ("\t", {"bold": True, "size": 16}),
            (heading, {"bold": True, "size": 16}),
        ],
        font=font,
    )
    if bookmark:
        add_bookmark(p, bookmark)
    return p


def _add_section_heading(doc: Document, roman: str, *, font: str, title: str = "", bookmark: str = "") -> None:
    _add_roman_heading(
        doc, roman, title or _SECTION_III, font=font, bookmark=bookmark
    )


def _emu(cm: float) -> int:
    return int(Cm(cm))


def _next_cover_drawing_id() -> int:
    global _COVER_DRAWING_ID
    _COVER_DRAWING_ID += 1
    return _COVER_DRAWING_ID


def _cover_year() -> str:
    return str(date.today().year)


def _xml_escape(text: str) -> str:
    return (
        (text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _float_inline_picture(inline, *, left_cm: float, top_cm: float, behind: bool = True) -> None:
    drawing = inline.getparent()
    extent = inline.find(qn("wp:extent"))
    cx = extent.get("cx") if extent is not None else str(_emu(21.0))
    cy = extent.get("cy") if extent is not None else str(_emu(29.7))
    doc_pr = inline.find(qn("wp:docPr"))
    graphic = inline.find(qn("a:graphic"))
    cnv = inline.find(qn("wp:cNvGraphicFramePr"))

    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "0")
    anchor.set("behindDoc", "1" if behind else "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    anchor.append(simple)

    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    off_h = OxmlElement("wp:posOffset")
    off_h.text = str(_emu(left_cm))
    pos_h.append(off_h)
    anchor.append(pos_h)

    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    off_v = OxmlElement("wp:posOffset")
    off_v.text = str(_emu(top_cm))
    pos_v.append(off_v)
    anchor.append(pos_v)

    new_extent = OxmlElement("wp:extent")
    new_extent.set("cx", cx)
    new_extent.set("cy", cy)
    anchor.append(new_extent)

    effect = OxmlElement("wp:effectExtent")
    for edge in ("l", "t", "r", "b"):
        effect.set(edge, "0")
    anchor.append(effect)
    anchor.append(OxmlElement("wp:wrapNone"))
    if doc_pr is not None:
        anchor.append(doc_pr)
    if cnv is not None:
        anchor.append(cnv)
    if graphic is not None:
        anchor.append(graphic)
    drawing.replace(inline, anchor)


def _add_cover_textbox(
    paragraph,
    text: str,
    *,
    left_cm: float,
    top_cm: float,
    width_cm: float,
    height_cm: float,
    font: str,
    size_pt: int,
    bold: bool = True,
    align: str = "left",
) -> None:
    shape_id = _next_cover_drawing_id()
    cx = _emu(width_cm)
    cy = _emu(height_cm)
    bold_xml = "<w:b/><w:bCs/>" if bold else ""
    xml = (
        '<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0"'
        f' relativeHeight="251659264" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{_emu(left_cm)}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{_emu(top_cm)}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        f'<wp:docPr id="{shape_id}" name="Cover {shape_id}"/>'
        '<wp:cNvGraphicFramePr/>'
        "<a:graphic>"
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<wps:wsp><wps:cNvSpPr txBox=\"1\"/><wps:spPr>"
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom><a:noFill/><a:ln><a:noFill/></a:ln>'
        "</wps:spPr><wps:txbx><w:txbxContent><w:p>"
        f'<w:pPr><w:jc w:val="{align}"/><w:spacing w:before="0" w:after="0"/></w:pPr>'
        "<w:r><w:rPr>"
        f'<w:rFonts w:ascii="{_xml_escape(font)}" w:hAnsi="{_xml_escape(font)}"'
        f' w:cs="{_xml_escape(font)}"/>{bold_xml}'
        f'<w:sz w:val="{size_pt * 2}"/><w:szCs w:val="{size_pt * 2}"/>'
        "</w:rPr>"
        f"<w:t>{_xml_escape(text)}</w:t>"
        "</w:r></w:p></w:txbxContent></wps:txbx>"
        '<wps:bodyPr wrap="square" lIns="0" tIns="0" rIns="0" bIns="0" anchor="ctr"/>'
        "</wps:wsp></a:graphicData></a:graphic></wp:anchor></w:drawing>"
    )
    run = paragraph.add_run()
    run._r.append(parse_xml(xml))


def _write_title_page(
    doc: Document,
    *,
    inspection_name: str,
    font: str,
) -> None:
    paragraph = doc.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not _COVER_IMAGE.exists():
        raise FileNotFoundError(f"Нет подложки титула: {_COVER_IMAGE}")
    picture_run = paragraph.add_run()
    inline = picture_run.add_picture(str(_COVER_IMAGE), width=Cm(21.0), height=Cm(29.7))
    _float_inline_picture(inline._inline, left_cm=0.0, top_cm=0.0, behind=True)

    name = (inspection_name or "Проверка").strip()
    _add_cover_textbox(
        paragraph,
        _COVER_TITLE,
        left_cm=3.5,
        top_cm=6.2,
        width_cm=14.0,
        height_cm=1.8,
        font=font,
        size_pt=20,
        align="center",
    )
    _add_cover_textbox(
        paragraph,
        name,
        left_cm=3.5,
        top_cm=9.4,
        width_cm=14.0,
        height_cm=4.5,
        font=font,
        size_pt=16,
        align="center",
    )
    _add_cover_textbox(
        paragraph,
        f"Минск {_cover_year()}",
        left_cm=3.0,
        top_cm=26.6,
        width_cm=5.8,
        height_cm=1.3,
        font=font,
        size_pt=14,
        align="left",
    )
    _add_cover_textbox(
        paragraph,
        _COVER_DEPT,
        left_cm=9.0,
        top_cm=26.6,
        width_cm=10.2,
        height_cm=1.3,
        font=font,
        size_pt=14,
        align="right",
    )
    break_run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    break_run._r.append(br)


def _enable_update_fields(doc: Document) -> None:
    settings_el = doc.settings.element
    if settings_el.find(qn("w:updateFields")) is not None:
        return
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings_el.append(el)


def _add_page_field(paragraph, *, font: str, size: int = 9) -> None:
    def fld(kind: str) -> None:
        run = paragraph.add_run()
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), kind)
        run._r.append(el)
        _set_run_font(run, size=size, font=font)

    fld("begin")
    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)
    _set_run_font(instr_run, size=size, font=font)
    fld("separate")
    placeholder = paragraph.add_run("2")
    _set_run_font(placeholder, size=size, font=font)
    fld("end")


def _write_toc_line(
    doc: Document,
    label: str,
    title: str,
    *,
    font: str,
    page: int | None,
    indent_cm: float = 0.0,
    bold_label: bool = True,
) -> None:
    paragraph = _add_styled_paragraph(
        doc, font=font, align="left", space_after=2, first_line=False
    )
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.left_indent = Cm(indent_cm)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Cm(16.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    heading = _dot_title(title)
    if bold_label:
        _add_runs(
            paragraph,
            [(f"{label}.", {"bold": True}), ("  ", {}), (heading, {})],
            font=font,
        )
    else:
        _add_runs(paragraph, [(f"{label}.  {heading}", {})], font=font)
    paragraph.add_run("\t")
    number = "—" if page is None else str(page)
    run = paragraph.add_run(number)
    _set_run_font(run, size=_FONT_SIZE, font=font)


def _write_toc(
    doc: Document,
    report: ConclusionDocument,
    *,
    font: str,
    opinion_body: str = "",
) -> None:
    heading = _add_styled_paragraph(
        doc, font=font, align="left", space_after=14, first_line=False
    )
    _add_runs(heading, [(_TOC_HEADING, {"bold": True, "size": 16})], font=font)
    iii_title = _SECTION_III
    observations: list[Observation] = []
    for section in report.sections:
        if section.kind == "observations":
            if section.title:
                iii_title = section.title
            observations = list(section.observations)
            break
    pages = estimate_toc_pages(opinion_body, report)
    _write_toc_line(doc, "I", _SECTION_I, font=font, page=pages.get("I", 3))
    _write_toc_line(doc, "II", _SECTION_II, font=font, page=None)
    _write_toc_line(doc, "III", iii_title, font=font, page=pages.get("III", pages.get("I", 3) + 1))
    for observation in observations:
        _write_toc_line(
            doc,
            observation.number,
            observation.title or "Наблюдение",
            font=font,
            page=pages.get(observation.number, pages.get("III")),
            indent_cm=1.0,
            bold_label=False,
        )
    _write_toc_line(doc, "IV", _SECTION_LAST, font=font, page=pages.get("IV"))
    _blank(doc, font=font, after=0)
    doc.add_page_break()


def _write_general_section(doc: Document, section: ReportSection, *, font: str) -> None:
    _add_roman_heading(doc, section.roman, _SECTION_LAST, font=font, bookmark="sec_IV")
    for label, value in section.general_items:
        lab = _add_styled_paragraph(
            doc, font=font, align="justify", space_before=8, space_after=2, first_line=False
        )
        _add_runs(lab, [(label, {"bold": True})], font=font)
        _add_opinion_paragraph(
            doc,
            value or "уточняется",
            font=font,
            size=_FONT_SIZE,
            first_line=False,
            space_after=8,
        )


def write_conclusion_docx(
    path: Path,
    *,
    inspection_name: str,
    case_id: str,
    opinion_body: str,
    report: ConclusionDocument,
    font: str = "Times New Roman",
) -> Path:
    doc = Document()
    _set_document_base_font(doc, font, _FONT_SIZE)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    section.different_first_page_header_footer = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer.paragraph_format.tab_stops.add_tab_stop(
        Cm(16.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
    )
    fr = footer.add_run(f"Черновик · {_TITLE} · {inspection_name} · кейс {case_id}")
    _set_run_font(fr, size=9, italic=True, font=font)
    footer.add_run("\t")
    _add_page_field(footer, font=font, size=9)

    _enable_update_fields(doc)
    _write_title_page(doc, inspection_name=inspection_name, font=font)
    _write_toc(doc, report, font=font, opinion_body=opinion_body or "")

    _add_roman_heading(doc, "I", _SECTION_I, font=font, bookmark="sec_I")
    add_opinion_markdown(doc, opinion_body or "", font=font)
    doc.add_page_break()

    for block in report.sections:
        if block.kind == "general":
            _write_general_section(doc, block, font=font)
            continue
        _add_section_heading(
            doc, block.roman, font=font, title=block.title or _SECTION_III, bookmark="sec_III"
        )
        if block.intro:
            _add_body_markdown(doc, block.intro, font=font)
        for observation in block.observations:
            _add_observation_header(doc, observation, font=font)
            if observation.body:
                _add_body_markdown(doc, observation.body, font=font)
            _add_observation_summary(doc, observation, font=font)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
