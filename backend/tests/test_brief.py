import tempfile
import unittest
from unittest.mock import patch
import zipfile
from pathlib import Path

from app.models import CaseState, KnowledgeItem, ProposedDocument
from app.services.brief_docx import write_brief_docx
from app.services.brief_flow import collect_brief_sources
from app.services.citations import excerpt_for_cite, extract_article_ref, origin_url, pages_estimate
from app.services.knowledge_flow import _fragments_from_item


class TestArticleRef(unittest.TestCase):
    def test_article_line(self):
        text = "Статья 625. Договор аренды\nАрендодатель обязуется предоставить имущество."
        self.assertIn("625", extract_article_ref(text) or "")

    def test_st_abbrev(self):
        self.assertEqual(extract_article_ref("Ст. 12. Валютные операции"), "Ст. 12. Валютные операции")

    def test_punkt_fallback(self):
        self.assertEqual(extract_article_ref("согласно пункт 3.2 инструкции"), "пункт 3.2")

    def test_none_when_empty(self):
        self.assertIsNone(extract_article_ref("просто абзац без номера"))


class TestExcerpt(unittest.TestCase):
    def test_trims_and_ellipsis(self):
        blob = "слово " * 200
        out = excerpt_for_cite(blob, max_chars=40)
        self.assertTrue(out.endswith("…"))
        self.assertLessEqual(len(out), 40)


class TestOriginUrl(unittest.TestCase):
    def test_matches_origin_document_id(self):
        doc = ProposedDocument(
            id="abc123abc123",
            title="ГК РБ",
            doc_type="кодекс",
            why_needed="x",
            found_url="https://pravo.by/document/?guid=3871&p0=hk9800218",
        )
        item = KnowledgeItem(
            title="ГК РБ",
            filename="gk.txt",
            local_path="/tmp/gk.txt",
            origin_document_id="abc123abc123",
        )
        state = CaseState(case_id="c1", inspection_name="Аренда", documents=[doc], knowledge=[item])
        self.assertIn("pravo.by", origin_url(state, item) or "")


class TestPages(unittest.TestCase):
    def test_six_pages(self):
        text = "а" * 11000
        self.assertGreaterEqual(pages_estimate(text, 1800), 6)


class TestBriefDocx(unittest.TestCase):
    def test_writes_hyperlinks_and_bookmarks(self):
        sources = [
            {
                "n": 1,
                "title": "Гражданский кодекс Республики Беларусь",
                "article": "Статья 625. Договор аренды",
                "excerpt": "По договору аренды арендодатель обязуется предоставить имущество.",
                "url": "https://pravo.by/document/?guid=3871&p0=hk9800218",
                "filename": "gk.txt",
            }
        ]
        chapters = [
            {
                "title": "Гражданский кодекс Республики Беларусь",
                "body": "## Ключевые нормы\nДоговор аренды — [1].",
            }
        ]
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "brief.docx"
            write_brief_docx(
                path,
                inspection_name="Проверка аренды",
                period="2025",
                keywords=["аренда", "НДС"],
                case_id="3a23fb6db4a9",
                overview="Смотрите аренду в [1].",
                chapters=chapters,
                sources=sources,
            )
            self.assertTrue(path.exists())
            from docx import Document

            doc = Document(str(path))
            texts = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Саммари нормативной базы", texts)
            self.assertIn("Статья 625", texts)
            self.assertIn("[1]", texts)
            with zipfile.ZipFile(path) as zf:
                xml = zf.read("word/document.xml").decode("utf-8")
                rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
            self.assertIn("cite_1", xml)
            self.assertIn("w:hyperlink", xml)
            self.assertIn("pravo.by/document/?guid=3871", rels)
            self.assertIn("TargetMode=\"External\"", rels)


class TestCollectSources(unittest.TestCase):
    def test_numbers_globally_and_keeps_article(self):
        with tempfile.TemporaryDirectory() as tmp:
            txt = Path(tmp) / "gk.txt"
            txt.write_text(
                "Статья 625. Договор аренды\n"
                "Арендодатель предоставляет имущество за плату.\n\n"
                "Статья 587. Существенные условия\n"
                "В договоре указывают объект аренды.\n",
                encoding="utf-8",
            )
            item = KnowledgeItem(
                id="item1",
                title="ГК РБ",
                filename="gk.txt",
                local_path=str(txt),
                text_path=str(txt),
                extract_status="ok",
            )
            doc = ProposedDocument(
                title="ГК РБ",
                doc_type="кодекс",
                why_needed="аренда",
                found_url="https://pravo.by/document/?guid=3871&p0=hk9800218",
            )
            state = CaseState(
                case_id="c1",
                inspection_name="Проверка аренды коммерческой недвижимости",
                keywords=["аренда"],
                documents=[doc],
                knowledge=[item],
            )
            sources = collect_brief_sources(state)
            self.assertTrue(sources)
            self.assertEqual(sources[0]["n"], 1)
            self.assertTrue(any(s.get("article") for s in sources))
            self.assertTrue(all(s.get("url") for s in sources))
            frags = _fragments_from_item(state, item)
            self.assertTrue(frags)

    def test_covers_start_and_end_not_keyword_hits(self):
        with tempfile.TemporaryDirectory() as tmp:
            blocks = []
            for i in range(1, 25):
                payload = "валюта " if i == 12 else "общее "
                blocks.append(f"Статья {i}. Раздел {i}\n{payload * 220}\n")
            txt = Path(tmp) / "code.txt"
            txt.write_text("\n".join(blocks), encoding="utf-8")
            item = KnowledgeItem(
                id="item1",
                title="Кодекс",
                filename="code.txt",
                text_path=str(txt),
                local_path=str(txt),
                extract_status="ok",
            )
            state = CaseState(
                case_id="c1",
                inspection_name="Проверка аренды",
                keywords=["валюта"],
                knowledge=[item],
            )
            blob = " ".join(fr["text"] for fr in _fragments_from_item(state, item))
            self.assertIn("Статья 1.", blob)
            self.assertIn("Статья 24.", blob)
            sources = collect_brief_sources(state)
            joined = " ".join(s["text"] for s in sources)
            self.assertIn("Статья 1.", joined)
            self.assertIn("Статья 24.", joined)


class TestSequentialCoverage(unittest.TestCase):
    def test_even_sample_keeps_first_and_last(self):
        from app.services.chunker import even_sample

        items = list(range(20))
        sampled = even_sample(items, 5)
        self.assertEqual(sampled[0], 0)
        self.assertEqual(sampled[-1], 19)
        self.assertEqual(len(sampled), 5)

    def test_windows_cover_every_article(self):
        from app.services.chunker import sequential_windows

        text = "\n".join(
            f"Статья {i}. Норма {i}\n{'слово ' * 180}" for i in range(1, 16)
        )
        windows = sequential_windows(text, size=2500, overlap=200)
        self.assertGreaterEqual(len(windows), 2)
        self.assertIn("Статья 1.", windows[0])
        self.assertIn("Статья 15.", windows[-1])
        for i in range(1, 16):
            self.assertTrue(
                any(f"Статья {i}." in w for w in windows),
                f"статья {i} выпала из последовательного чтения",
            )

    def test_windows_keep_article_boundary(self):
        from app.services.chunker import sequential_windows

        text = (
            "Статья 1. Первая\n" + ("ааа " * 400) + "\n"
            "Статья 2. Вторая\n" + ("ббб " * 400) + "\n"
        )
        windows = sequential_windows(text, size=2500, overlap=0)
        self.assertGreaterEqual(len(windows), 2)
        self.assertTrue(windows[0].startswith("Статья 1."))
        self.assertTrue(any(w.startswith("Статья 2.") for w in windows))

    def test_article_outline_order(self):
        from app.services.citations import extract_article_outline

        text = (
            "Глава 1. Общие положения\n"
            "Статья 1. Предмет\nтекст\n"
            "Статья 2. Термины\nтекст\n"
        )
        outline = extract_article_outline(text)
        self.assertEqual(outline[0], "Глава 1. Общие положения")
        self.assertIn("Статья 1. Предмет", outline)
        self.assertIn("Статья 2. Термины", outline)


class TestDownloadNames(unittest.TestCase):
    def test_summary_docx_uses_inspection_name(self):
        from app.services.brief_flow import brief_download_name

        name = brief_download_name("Проверка аренды коммерческой недвижимости", "abc")
        self.assertTrue(name.endswith("_summary.docx"))
        self.assertIn("аренды", name)
        self.assertNotIn("abc", name)

    def test_total_docx_uses_inspection_name(self):
        from app.services.total_flow import total_download_name

        name = total_download_name("Проверка аренды коммерческой недвижимости", "abc")
        self.assertTrue(name.endswith("_total.docx"))
        self.assertIn("аренды", name)
        self.assertNotIn("abc", name)

    def test_archive_zip_uses_npa_suffix(self):
        from app.storage import CaseStore

        name = CaseStore.archive_filename("Проверка аренды коммерческой недвижимости", "abc")
        self.assertTrue(name.endswith("_npa.zip"))
        self.assertNotIn("abc", name)


class TestTotalParse(unittest.TestCase):
    def test_parse_sources_section(self):
        from app.services.total_flow import parse_total_sources

        md = """## Суть темы
Аренда важна [1].

## Источники
[1] Гражданский кодекс РБ — Статья 625 — https://pravo.by/document/?guid=3871
[2] Банковский кодекс РБ — глава о договорах — URL неизвестен
"""
        body, sources = parse_total_sources(md)
        self.assertIn("Аренда важна", body)
        self.assertNotIn("## Источники", body)
        self.assertEqual(len(sources), 2)
        self.assertEqual(sources[0]["n"], 1)
        self.assertIn("Гражданский", sources[0]["title"])
        self.assertEqual(sources[0]["article"], "Статья 625")
        self.assertIn("pravo.by", sources[0]["url"])
        self.assertEqual(sources[1]["url"], "")

    def test_writes_total_docx(self):
        from app.services.brief_docx import write_total_docx

        sources = [
            {
                "n": 1,
                "title": "Гражданский кодекс Республики Беларусь",
                "article": "Статья 625",
                "url": "https://pravo.by/document/?guid=3871&p0=hk9800218",
                "excerpt": "",
            }
        ]
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "total.docx"
            write_total_docx(
                path,
                inspection_name="Проверка аренды",
                period="2025",
                keywords=["аренда"],
                case_id="c1",
                body="## Суть темы\nДоговор аренды [1].",
                sources=sources,
            )
            self.assertTrue(path.exists())
            with zipfile.ZipFile(path) as zf:
                xml = zf.read("word/document.xml").decode("utf-8")
            self.assertIn("знания модели", xml)
            self.assertIn("cite_1", xml)


class TestProgramDocx(unittest.TestCase):
    def test_programma_name_uses_inspection(self):
        from app.services.program_flow import program_download_name

        name = program_download_name("Проверка аренды коммерческой недвижимости", "abc")
        self.assertTrue(name.endswith("_programma.docx"))
        self.assertIn("аренды", name)
        self.assertNotIn("abc", name)

    def test_writes_program_title_and_cites(self):
        from app.services.brief_docx import write_program_docx

        sources = [
            {
                "n": 1,
                "title": "Гражданский кодекс Республики Беларусь",
                "article": "Статья 625. Договор аренды",
                "excerpt": "По договору аренды арендодатель обязуется предоставить имущество.",
                "url": "https://pravo.by/document/?guid=3871&p0=hk9800218",
                "filename": "gk.txt",
            }
        ]
        questions = [
            "Анализ договоров аренды. Критерий: [1].",
            "Проверка полномочий при заключении договоров аренды.",
        ]
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "program.docx"
            write_program_docx(
                path,
                inspection_name="Проверка аренды коммерческой недвижимости",
                period="2025",
                keywords=["аренда", "НДС"],
                case_id="3a23fb6db4a9",
                body="## Вопросы, подлежащие аудиту\n\n1. Анализ договоров аренды. Критерий: [1].",
                sources=sources,
                questions=questions,
            )
            self.assertTrue(path.exists())
            from docx import Document

            doc = Document(str(path))
            texts = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("ПРОГРАММА", texts)
            self.assertGreaterEqual(len(doc.tables), 3)
            info, questions_table, sign_table = doc.tables[0], doc.tables[1], doc.tables[2]
            info_text = "\n".join(cell.text for row in info.rows for cell in row.cells)
            self.assertIn("Название проверки", info_text)
            self.assertIn("Проверка аренды коммерческой недвижимости", info_text)
            self.assertIn("Аудируемый период", info_text)
            self.assertIn("2025", info_text)
            header = " ".join(cell.text for cell in questions_table.rows[0].cells)
            self.assertIn("Вопросы, подлежащие аудиту", header)
            body_text = "\n".join(
                cell.text for row in questions_table.rows[1:] for cell in row.cells
            )
            self.assertIn("Анализ договоров аренды", body_text)
            self.assertIn("[1]", body_text)
            self.assertIn("2.", body_text)
            self.assertIn("При необходимости вопросы, подлежащие аудиту", texts)
            sign = "\n".join(cell.text for row in sign_table.rows for cell in row.cells)
            self.assertIn("Менеджер по направлению деятельности", sign)
            info_width = sum(int(col.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w")) for col in info._tbl.tblGrid)
            q_width = sum(int(col.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w")) for col in questions_table._tbl.tblGrid)
            self.assertEqual(info_width, q_width)


class TestProgramItems(unittest.TestCase):
    def test_normalize_default_and_range(self):
        from app.services.program_flow import normalize_program_item_range

        self.assertEqual(normalize_program_item_range(), (8, 11))
        self.assertEqual(normalize_program_item_range(items="8"), (8, 8))
        self.assertEqual(normalize_program_item_range(items="10-12"), (10, 12))
        self.assertEqual(normalize_program_item_range(items="5 - 6"), (5, 6))
        self.assertEqual(normalize_program_item_range(items_min=8, items_max=8), (8, 8))
        self.assertEqual(normalize_program_item_range(items="1"), (3, 3))
        self.assertEqual(normalize_program_item_range(items="99"), (20, 20))

    def test_parse_questions_numbered_list(self):
        from app.services.program_flow import parse_program_questions

        body = """
## Название проверки
Проверка аренды

## Вопросы, подлежащие аудиту

1. Анализ ЛПА банка по аренде. Критерий [1].
2. Проверка полномочий при заключении договоров.
3. Анализ арендных платежей и коммунальных расходов.
"""
        questions = parse_program_questions(body)
        self.assertEqual(len(questions), 3)
        self.assertIn("Анализ ЛПА", questions[0])
        self.assertIn("арендных платежей", questions[2])

    def test_parse_questions_skips_nested_numbering_and_clips(self):
        from app.services.program_flow import fit_program_questions, parse_program_questions

        body = """
## Вопросы, подлежащие аудиту

1. Первый вопрос. Внутри не отдельный пункт.
2. Второй вопрос.
3. Третий.
4. Четвёртый.
5. Пятый.
6. Шестой.
7. Седьмой.
8. Восьмой.
9. Девятый лишний.
10. Десятый лишний.
"""
        questions = parse_program_questions(body)
        self.assertEqual(len(questions), 10)
        clipped = fit_program_questions(questions, 8)
        self.assertEqual(len(clipped), 8)
        self.assertIn("Восьмой", clipped[-1])
        self.assertTrue(all("лишний" not in q for q in clipped))

    def test_program_prompts_require_three_to_four_sentences(self):
        from app.prompts import prompt
        from app.services.program_flow import program_items_hint

        sections = prompt("program_sections", items_hint=program_items_hint(8, 8))
        system = prompt("program_system")
        user = prompt(
            "program_user",
            inspection="Проверка аренды",
            keywords="аренда",
            period="2025",
            document_catalog="",
            catalog="",
            fragments="",
            cards_block="",
            sections=sections,
            items_hint=program_items_hint(8, 8),
            target=4800,
            target_hi=5760,
        )
        self.assertIn("3–4", sections)
        self.assertIn("3–4", system)
        self.assertIn("3–4", user)
        self.assertNotIn("~40 слов", sections)
        self.assertNotIn("1–2 коротких", sections)

    def test_parse_questions_keeps_multi_sentence_item(self):
        from app.services.program_flow import parse_program_questions

        body = """
## Вопросы, подлежащие аудиту

1. Проверка существенных условий договоров аренды: описание объекта, срок и размер платы. Сопоставить основной договор с допсоглашениями на согласованность предмета и расчётов. Запросить договоры, допсоглашения и приложения. Критерий [3].
2. Верификация полномочий арендодателя на сдачу имущества. Сверить правоустанавливающие документы с условиями договора. Запросить выписки ЕГР, доверенности, уставы. Критерий [3].
"""
        questions = parse_program_questions(body)
        self.assertEqual(len(questions), 2)
        self.assertIn("Сопоставить основной договор", questions[0])
        self.assertIn("Критерий [3]", questions[0])
        self.assertIn("правоустанавливающие документы", questions[1])

    def test_parse_questions_ignores_restarted_list(self):
        from app.services.program_flow import parse_program_questions

        body = """
## Вопросы, подлежащие аудиту

1. Анализ ЛПА. Запросить: 1. карту процесса 2. регламент.
2. Проверка полномочий.
"""
        questions = parse_program_questions(body)
        self.assertEqual(len(questions), 2)
        self.assertIn("карту процесса", questions[0])
        self.assertIn("полномочий", questions[1])


class TestSummarizeFullDocument(unittest.IsolatedAsyncioTestCase):
    async def test_oneshot_prompt_contains_whole_text(self):
        from app.services.knowledge_flow import summarize_item

        with tempfile.TemporaryDirectory() as tmp:
            txt = Path(tmp) / "act.txt"
            txt.write_text(
                "Статья 1. Предмет\nДоговор аренды недвижимого имущества.\n\n"
                "Статья 12. Расчёты\nОплата в белорусских рублях.\n\n"
                "Статья 40. Ответственность\nНеустойка за просрочку.\n",
                encoding="utf-8",
            )
            item = KnowledgeItem(
                title="Инструкция тестовая",
                filename="act.txt",
                text_path=str(txt),
                local_path=str(txt),
                extract_status="ok",
            )
            state = CaseState(
                case_id="c1",
                inspection_name="Проверка аренды",
                keywords=["валюта"],
                knowledge=[item],
            )
            prompts: list[str] = []

            async def fake_chat(system, user, **kwargs):
                prompts.append(user)
                return (
                    "## Зачем этот акт\nдля проверки аренды\n"
                    "## Ключевые нормы\n"
                    "- ст. 1 — аренда\n- ст. 12 — расчёты\n- ст. 40 — неустойка"
                )

            with patch("app.services.knowledge_flow.chat_complete", fake_chat):
                result = await summarize_item(state, item)

            self.assertEqual(result.summary_status, "ok")
            self.assertEqual(len(prompts), 1)
            self.assertIn("Статья 1. Предмет", prompts[0])
            self.assertIn("Статья 40. Ответственность", prompts[0])
            self.assertIn("целиком", prompts[0].lower())
            self.assertIn("Ключевые слова", prompts[0])
            self.assertIn("валюта", prompts[0])
            self.assertIn("Проверка аренды", prompts[0])
            self.assertIn("Зачем этот акт", prompts[0])

    async def test_long_document_uses_rag_not_map_reduce(self):
        from app.services.knowledge_flow import summarize_item

        with tempfile.TemporaryDirectory() as tmp:
            parts = []
            for i in range(1, 12):
                payload = ("аренда валюта расчёты " if i in (3, 7) else "общее положение ") * 200
                parts.append(f"Статья {i}. Норма {i}\n{payload}\n")
            txt = Path(tmp) / "code.txt"
            txt.write_text("\n".join(parts), encoding="utf-8")
            item = KnowledgeItem(
                id="itemrag1",
                title="Кодекс",
                filename="code.txt",
                text_path=str(txt),
                local_path=str(txt),
                extract_status="ok",
            )
            state = CaseState(
                case_id="c1",
                inspection_name="Проверка аренды",
                keywords=["аренда", "валюта"],
                knowledge=[item],
            )
            prompts: list[str] = []

            async def fake_chat(system, user, **kwargs):
                prompts.append(user)
                return (
                    "## Зачем этот акт\nКодекс задаёт рамку проверки.\n"
                    "## Суть и сфера\nГражданские договоры.\n"
                    "## Ключевые нормы\n- ст. 3 — аренда [1]\n- ст. 7 — валюта [2]\n"
                    "## Что проверять\n- предмет договора\n"
                    "## Чего нет в тексте\nв выборке нет раздела про налоги"
                )

            async def fake_embed(texts):
                # Deterministic tiny vectors: lease/currency keywords → axis 0
                out = []
                for t in texts:
                    low = t.lower()
                    out.append([1.0 if ("аренда" in low or "валюта" in low) else 0.0, 0.1])
                return out

            with (
                patch("app.services.knowledge_flow.chat_complete", fake_chat),
                patch("app.services.knowledge_flow.embed_texts", fake_embed),
                patch("app.services.knowledge_flow._load_index", return_value={"chunks": []}),
                patch("app.services.knowledge_flow._persist_item_embeddings"),
            ):
                result = await summarize_item(state, item)

            self.assertEqual(result.summary_status, "ok")
            self.assertEqual(len(prompts), 1)
            self.assertTrue(any("rag" in p.lower() or "выборк" in p.lower() for p in prompts))
            self.assertFalse(any("часть 1 из" in p.lower() for p in prompts))
            self.assertFalse(any("собери из заметок" in p.lower() for p in prompts))
            joined = prompts[0]
            self.assertIn("аренда", joined.lower())
            self.assertIn("Ключевые нормы", result.summary)
            self.assertTrue(result.citations)
            # RAG should prefer topical articles over the whole code dump
            cite_blob = " ".join(c.get("text") or "" for c in result.citations)
            self.assertTrue(
                "Статья 3." in cite_blob or "Статья 7." in cite_blob,
                "RAG citations should include keyword-relevant articles",
            )


class TestRetrieveEvidence(unittest.IsolatedAsyncioTestCase):
    async def test_mmr_and_bm25_prefer_query_hits(self):
        from app.services.knowledge_retrieve import select_evidence

        chunks = []
        for i in range(1, 15):
            payload = ("аренда валюта " if i in (4, 10) else "канцелярия ") * 40
            chunks.append(
                {
                    "id": f"doc:{i}",
                    "item_id": "doc",
                    "title": "Кодекс",
                    "filename": "c.txt",
                    "text": f"Статья {i}.\n{payload}",
                    "embedding": [],
                }
            )

        async def fake_embed(texts):
            out = []
            for t in texts:
                low = t.lower()
                out.append([1.0 if ("аренда" in low or "валюта" in low) else 0.0, 0.2])
            return out

        picked = await select_evidence(
            chunks,
            ["Проверка аренды", "аренда", "валюта"],
            top_k=4,
            candidates=20,
            neighbor=0,
            always_include_first=True,
            embed_fn=fake_embed,
        )
        blob = " ".join(p["text"] for p in picked)
        self.assertIn("Статья 4.", blob)
        self.assertIn("Статья 10.", blob)
        self.assertLessEqual(len(picked), 6)

class TestOverviewAndAskHelpers(unittest.TestCase):
    def test_overview_lists_every_act_without_llm(self):
        from app.services.brief_flow import _synthesize

        state = CaseState(case_id="c1", inspection_name="Проверка аренды")
        chapters = [
            {
                "title": "Инструкция № 1",
                "body": "## Ключевые нормы\n- ст. 1 — предмет\n- ст. 2 — сроки",
            },
            {
                "title": "Кодекс",
                "body": "## Ключевые нормы\n- ст. 625 — аренда\n- ст. 626 — плата",
            },
        ]
        overview = _synthesize(state, chapters)
        self.assertIn("Инструкция № 1", overview)
        self.assertIn("Кодекс", overview)
        self.assertIn("ст. 1 — предмет", overview)
        self.assertIn("ст. 625 — аренда", overview)


class TestAskRetrieval(unittest.IsolatedAsyncioTestCase):
    def test_article_query_variants(self):
        from app.services.knowledge_retrieve import article_query_variants, ask_queries

        variants = article_query_variants("Какой срок в ст. 625 ГК?")
        self.assertIn("Статья 625", variants)
        queries = ask_queries("ст. 625 аренда")
        self.assertTrue(any(q.startswith("Статья 625") for q in queries))

    async def test_heading_beats_cross_reference(self):
        from app.services.knowledge_retrieve import retrieve_for_ask

        chunks = [
            {
                "id": "instr:0",
                "item_id": "instr",
                "title": "Инструкция",
                "filename": "i.txt",
                "text": "Статья 12. Общие положения\nсм. также статью 625 Гражданского кодекса про аренду.",
                "embedding": [],
            },
            {
                "id": "gk:0",
                "item_id": "gk",
                "title": "Гражданский кодекс",
                "filename": "gk.txt",
                "text": "Статья 624. Преимущественное право\nканцелярия " * 20,
                "embedding": [],
            },
            {
                "id": "gk:1",
                "item_id": "gk",
                "title": "Гражданский кодекс",
                "filename": "gk.txt",
                "text": "Статья 625. Договор аренды\nсрок регистрации договора аренды недвижимости " * 20,
                "embedding": [],
            },
            {
                "id": "gk:2",
                "item_id": "gk",
                "title": "Гражданский кодекс",
                "filename": "gk.txt",
                "text": "Статья 626. Арендная плата\nканцелярия " * 20,
                "embedding": [],
            },
        ]

        async def fake_embed(texts):
            out = []
            for t in texts:
                low = t.lower()
                out.append([1.0 if "625" in low or "аренда" in low else 0.0, 0.2])
            return out

        picked = await retrieve_for_ask(chunks, "вопрос ст. 625 срок аренды", top_k=1, embed_fn=fake_embed)
        self.assertTrue(picked)
        self.assertIn("Статья 625.", picked[0]["text"])
        blob = " ".join(p["text"] for p in picked)
        self.assertNotIn("Статья 12.", blob)
        self.assertNotIn("Статья 624.", blob)
        self.assertNotIn("Статья 626.", blob)

    async def test_neighbor_does_not_glue_next_article(self):
        from app.services.knowledge_retrieve import select_evidence

        chunks = [
            {
                "id": "doc:0",
                "item_id": "doc",
                "title": "Кодекс",
                "filename": "c.txt",
                "text": "Статья 3. Аренда валюта расчёты\n" + ("аренда валюта " * 40),
                "embedding": [],
            },
            {
                "id": "doc:1",
                "item_id": "doc",
                "title": "Кодекс",
                "filename": "c.txt",
                "text": "Статья 4. Канцелярия\n" + ("канцелярия " * 40),
                "embedding": [],
            },
        ]

        async def fake_embed(texts):
            return [[1.0 if "аренда" in t.lower() else 0.0, 0.1] for t in texts]

        picked = await select_evidence(
            chunks,
            ["аренда"],
            top_k=1,
            candidates=4,
            neighbor=1,
            always_include_first=False,
            embed_fn=fake_embed,
        )
        blob = " ".join(p["text"] for p in picked)
        self.assertIn("Статья 3.", blob)
        self.assertNotIn("Статья 4.", blob)


if __name__ == "__main__":
    unittest.main()
