import re
import tempfile
import unittest
import zipfile
from datetime import date
from pathlib import Path

from app.prompts import prompt
from app.services.conclusion_docx import (
    _COVER_IMAGE,
    _cover_year,
    ensure_all_hypotheses,
    estimate_toc_pages,
    parse_conclusion_markdown,
    toc_entries,
    write_conclusion_docx,
)
from app.services.opinion_flow import FONT_CALIBRI, FONT_TIMES, parse_opinion_font_flag


def _row(n: int, priority: str = "средний") -> dict:
    return {
        "n": str(n),
        "hypothesis": f"Гипотеза {n} о риске учёта",
        "priority": priority,
        "assertion": "полнота",
        "risk": "риск недостоверности",
        "why_risk": "нет сверки",
        "how_to_test": "сверка договоров",
        "evidence_request": "договоры",
        "working_paper": "Доработать сверку регистров с договорами.",
        "npa_criteria": "НК РБ",
        "plan_sections": "учёт",
    }


SAMPLE_MD = """
## Содержание
III. Учёт аренды и сверка договоров
IV. Общая информация об аудиторской проверке

## Раздел III. Учёт аренды и сверка договоров

В ходе проверки рассмотрены вопросы отражения арендных платежей.

### Наблюдение 3.1. Неполная сверка договоров аренды
существенность: высокий
гипотеза: 1

По подтверждённой гипотезе отмечается риск неполного отражения платежей по договорам аренды.
Контроль сверки регистров с договорами требует усиления.

рекомендация:
Обеспечить регулярную сверку регистров аренды с действующими договорами.

### Наблюдение 3.2. Своевременность отражения НДС
существенность: средний
гипотеза: 3

Отмечается риск несвоевременного отражения НДС по арендным операциям.

рекомендация:
Актуализировать контрольные процедуры по срокам отражения НДС.

## Раздел IV. Общая информация об аудиторской проверке

### Аудируемый период
2025

### Вид аудита
Тематическая аудиторская проверка
"""


class TestConclusionFont(unittest.TestCase):
    def test_parse_flags(self):
        self.assertEqual(parse_opinion_font_flag("аудиторское заключение -c"), FONT_CALIBRI)
        self.assertEqual(parse_opinion_font_flag("аудиторское заключение -t заново"), FONT_TIMES)
        self.assertEqual(parse_opinion_font_flag("аудиторское заключение"), FONT_TIMES)


class TestConclusionParse(unittest.TestCase):
    def test_parses_observations_and_skips_i_ii(self):
        report = parse_conclusion_markdown(
            SAMPLE_MD,
            hypotheses=[_row(1, "высокий"), _row(3, "средний")],
            period="2025",
        )
        obs_sections = [s for s in report.sections if s.kind == "observations"]
        general = [s for s in report.sections if s.kind == "general"]
        self.assertEqual(len(obs_sections), 1)
        self.assertEqual(len(obs_sections[0].observations), 2)
        self.assertEqual(obs_sections[0].observations[0].number, "3.1")
        self.assertEqual(obs_sections[0].observations[0].materiality, "высокий")
        self.assertNotIn("существенность", obs_sections[0].observations[0].body.lower())
        self.assertIn("сверк", obs_sections[0].observations[0].recommendation.lower())
        self.assertEqual(len(general), 1)
        self.assertEqual(general[0].roman, "IV")
        self.assertEqual(obs_sections[0].roman, "III")
        self.assertIn("учёт аренды", obs_sections[0].title.lower())
        titles = [title for _, title in toc_entries(report)]
        self.assertEqual(
            titles,
            [
                "Аудиторское мнение по итогам проверки.",
                "Основные результаты аудита и итоговые аудиторские рекомендации.",
                "Учёт аренды и сверка договоров.",
                "Общая информация об аудиторской проверке.",
            ],
        )
        self.assertEqual([roman for roman, _ in toc_entries(report)], ["I", "II", "III", "IV"])

    def test_merges_extra_observation_sections_into_iii(self):
        md = """
## Содержание
III. Тема А
IV. Тема Б
V. Общая информация об аудиторской проверке

## Раздел III. Тема А
### Наблюдение 3.1. Первое
существенность: высокий
гипотеза: 1
Текст первого наблюдения.
рекомендация:
Сделать первое.

## Раздел IV. Тема Б
### Наблюдение 4.1. Второе
существенность: средний
гипотеза: 3
Текст второго наблюдения.
рекомендация:
Сделать второе.

## Раздел V. Общая информация об аудиторской проверке
### Аудируемый период
2025
"""
        report = parse_conclusion_markdown(
            md,
            hypotheses=[_row(1, "высокий"), _row(3, "средний")],
            period="2025",
        )
        obs_sections = [s for s in report.sections if s.kind == "observations"]
        self.assertEqual(len(obs_sections), 1)
        self.assertEqual(len(obs_sections[0].observations), 2)
        self.assertEqual(obs_sections[0].observations[0].number, "3.1")
        self.assertEqual(obs_sections[0].observations[1].number, "3.2")
        self.assertEqual(
            [title for _, title in toc_entries(report)][2],
            "Тема А.",
        )

    def test_fallback_uses_all_hypotheses(self):
        report = parse_conclusion_markdown(
            "просто текст без структуры",
            hypotheses=[_row(1, "высокий"), _row(2, "низкий")],
            period="2024",
        )
        obs = report.sections[0].observations
        self.assertEqual(len(obs), 2)
        self.assertEqual(obs[0].materiality, "высокий")
        self.assertEqual(obs[1].materiality, "низкий")

    def test_toc_page_numbers_increase_with_observations(self):
        rows = [_row(n) for n in range(1, 7)]
        report = parse_conclusion_markdown(
            "просто текст без структуры",
            hypotheses=rows,
            period="2025",
        )
        for obs in report.sections[0].observations:
            obs.body = ("Абзац наблюдения о риске контроля. " * 40).strip()
            obs.recommendation = "Закрепить процедуру сверки в ЛПА банка."
        pages = estimate_toc_pages("Цели аудита. " * 80, report)
        self.assertGreaterEqual(pages["I"], 3)
        self.assertGreater(pages["III"], pages["I"])
        self.assertGreaterEqual(pages["3.1"], pages["III"])
        self.assertGreater(pages["3.6"], pages["3.1"])
        self.assertGreaterEqual(pages["IV"], pages["3.6"])
        numbers = [pages["I"], pages["III"], pages["3.1"], pages["3.6"], pages["IV"]]
        self.assertGreater(len(set(numbers)), 1)

    def test_ensure_all_hypotheses_appends_missing(self):
        rows = [_row(n) for n in (1, 3, 4, 5, 6, 8)]
        report = parse_conclusion_markdown(SAMPLE_MD, hypotheses=rows[:2], period="2025")
        self.assertEqual(len(report.sections[0].observations), 2)
        filled = ensure_all_hypotheses(
            report,
            rows,
            period="2025",
            inspection_name="Проверка аренды коммерческой недвижимости",
        )
        self.assertEqual(len(filled.sections[0].observations), 6)
        self.assertEqual(
            [item.hypothesis_n for item in filled.sections[0].observations],
            ["1", "3", "4", "5", "6", "8"],
        )

    def test_tax_template_title_replaced_when_inspection_is_not_tax(self):
        md = """
## Раздел III. Оценка соответствия деятельности принципам налогообложения и защиты прав плательщиков.
### Наблюдение 3.1. Сверка договоров
существенность: высокий
гипотеза: 1
Текст.
рекомендация:
Сделать.
## Раздел IV. Общая информация об аудиторской проверке
### Аудируемый период
2025
"""
        report = parse_conclusion_markdown(
            md,
            hypotheses=[_row(1, "высокий")],
            period="2025",
            inspection_name="Проверка аренды коммерческой недвижимости",
        )
        self.assertIn("наблюдения по итогам", report.sections[0].title.lower())
        self.assertNotIn("налогообложения", report.sections[0].title.lower())


class TestConclusionDocx(unittest.TestCase):
    def test_cover_year_and_background_asset(self):
        self.assertTrue(_COVER_IMAGE.exists(), _COVER_IMAGE)
        self.assertEqual(_cover_year("2025"), "2025")
        self.assertEqual(_cover_year("2024 год и 3 квартала 2025 года"), "2025")

    def test_structure_font_and_observation_box(self):
        report = parse_conclusion_markdown(
            SAMPLE_MD,
            hypotheses=[_row(1, "высокий"), _row(3, "средний")],
            period="2025",
        )
        opinion = "## Аудиторское мнение\nПо подтверждённым гипотезам отмечается риск учёта аренды."
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "zakluchenie.docx"
            write_conclusion_docx(
                path,
                inspection_name="Проверка аренды коммерческой недвижимости",
                period="2025",
                case_id="abc123",
                opinion_body=opinion,
                report=report,
                font=FONT_CALIBRI,
            )
            self.assertTrue(path.exists())
            from docx import Document

            doc = Document(str(path))
            texts = "\n".join(p.text for p in doc.paragraphs)
            with zipfile.ZipFile(path) as zf:
                xml = zf.read("word/document.xml").decode("utf-8")
                media = [n for n in zf.namelist() if n.startswith("word/media/")]
            self.assertIn("Аудиторское заключение", xml)
            self.assertIn("Проверка аренды коммерческой недвижимости", xml)
            self.assertIn("Минск 2025", xml)
            self.assertIn("Департамент внутреннего аудита", xml)
            self.assertIn("Содержание", texts)
            self.assertIn("I.  Аудиторское мнение по итогам проверки.", texts)
            self.assertIn("II.  Основные результаты аудита и итоговые аудиторские рекомендации.", texts)
            self.assertIn("III.  Учёт аренды и сверка договоров.", texts)
            self.assertIn("3.1.  Неполная сверка договоров аренды.", texts)
            self.assertIn("IV.  Общая информация об аудиторской проверке.", texts)
            self.assertNotIn("Раздел III.", texts)
            self.assertNotIn("Разделы аудиторского заключения", texts)
            self.assertIn("Уровень существенности:", texts)
            self.assertIn("высокий", texts)
            self.assertIn("Аудитор:", texts)
            self.assertIn("Объект аудита:", texts)
            self.assertIn("Аудиторская рекомендация:", texts)
            self.assertIn("Срок –", texts)
            self.assertIn("Общая информация об аудиторской проверке.", texts)
            self.assertIn("По подтверждённым гипотезам отмечается риск учёта аренды.", texts)
            self.assertNotIn("Основные нарушения и недостатки", texts)
            boxes = [t for t in doc.tables if "Наблюдение" in t.cell(0, 0).text]
            self.assertGreaterEqual(len(boxes), 2)
            self.assertIn("3.1", boxes[0].cell(0, 0).text)
            self.assertTrue(boxes[0].cell(0, 1).paragraphs[0].runs[0].italic)
            from docx.oxml.ns import qn

            tbl_borders = boxes[0]._tbl.tblPr.find(qn("w:tblBorders"))
            self.assertIsNotNone(tbl_borders)
            self.assertEqual(tbl_borders.find(qn("w:top")).get(qn("w:val")), "nil")
            self.assertEqual(tbl_borders.find(qn("w:left")).get(qn("w:val")), "nil")
            boxed = {
                p.text.strip(): p._p.find(qn("w:pPr")).find(qn("w:pBdr")) is not None
                for p in doc.paragraphs
                if p._p.find(qn("w:pPr")) is not None
                and p._p.find(qn("w:pPr")).find(qn("w:pBdr")) is not None
            }
            boxed_text = "\n".join(boxed)
            self.assertIn("Уровень существенности:", boxed_text)
            self.assertIn("Аудитор:", boxed_text)
            self.assertIn("Объект аудита:", boxed_text)
            self.assertIn("Руководитель объекта аудита:", boxed_text)
            self.assertIn("Аудиторская рекомендация:", boxed_text)
            self.assertIn("Обеспечить регулярную сверку регистров аренды с действующими договорами.", boxed_text)
            self.assertIn("Срок –", boxed_text)
            self.assertIn("Calibri", xml)
            self.assertIn("w:tbl", xml)
            self.assertIn("w:pBdr", xml)
            self.assertIn('behindDoc="1"', xml)
            self.assertTrue(any(n.endswith(".png") for n in media), media)
            self.assertTrue(doc.sections[0].different_first_page_header_footer)
            self.assertLess(xml.find("Наблюдение 3.1"), xml.find("По подтверждённой гипотезе"))
            self.assertLess(
                xml.find("По подтверждённой гипотезе"),
                xml.find("Аудиторская рекомендация"),
            )
            self.assertNotIn("PAGEREF", xml)
            with zipfile.ZipFile(path) as footers:
                footer_xml = "\n".join(
                    footers.read(n).decode("utf-8")
                    for n in footers.namelist()
                    if n.startswith("word/footer")
                )
            self.assertIn("PAGE", footer_xml)
            toc_i = re.search(r"Аудиторское мнение по итогам проверки\.\t(\d+)", texts)
            toc_iv = re.search(r"Общая информация об аудиторской проверке\.\t(\d+)", texts)
            self.assertIsNotNone(toc_i)
            self.assertIsNotNone(toc_iv)
            self.assertGreaterEqual(int(toc_i.group(1)), 3)
            self.assertGreater(int(toc_iv.group(1)), int(toc_i.group(1)))
            self.assertIn("3.1.  Неполная сверка договоров аренды.", texts)

    def test_markdown_table_and_scheme_in_observation(self):
        md = """
## Раздел III. Наблюдения по итогам проверки
### Наблюдение 3.1. Контроль договоров
существенность: высокий
гипотеза: 1
Проверяется учёт аренды.

| Требование | Риск |
|---|---|
| Сверка договоров [1] | Пропуск объекта |

Схема: договор → регистр → платёж

рекомендация:
Ввести ежемесячную сверку.
## Раздел IV. Общая информация об аудиторской проверке
### Аудируемый период
2025
"""
        report = parse_conclusion_markdown(
            md,
            hypotheses=[_row(1, "высокий")],
            period="2025",
            inspection_name="Проверка аренды",
        )
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "zakluchenie.docx"
            write_conclusion_docx(
                path,
                inspection_name="Проверка аренды",
                period="2025",
                case_id="abc123",
                opinion_body="Мнение.",
                report=report,
                font=FONT_CALIBRI,
            )
            from docx import Document

            doc = Document(str(path))
            texts = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Схема: договор", texts)
            tables = [t for t in doc.tables if "Требование" in t.cell(0, 0).text]
            self.assertTrue(tables)
            self.assertIn("Сверка договоров [1]", tables[0].cell(1, 0).text)


class TestConclusionPrompts(unittest.TestCase):
    def test_prompts_skip_section_ii_and_require_template(self):
        system = prompt("conclusion_system")
        sections = prompt(
            "conclusion_sections",
            section_iii_title="Наблюдения по итогам проверки",
            hypothesis_count=6,
            hypothesis_numbers="1, 3, 4, 5, 6, 8",
            first_hypothesis_n="1",
        )
        user = prompt(
            "conclusion_user",
            inspection="Проверка аренды",
            keywords="аренда",
            period="2025",
            hypothesis_count=6,
            hypothesis_numbers="1, 3, 4, 5, 6, 8",
            observation_outline="3.1 ← гипотеза 1: тест",
            document_catalog="",
            hypotheses_block="1. тест",
            opinion_block="",
            program_block="",
            brief_block="",
            total_block="",
            cards_block="",
            fragments="",
            sections=sections,
        )
        self.assertIn("подтверждённ", system)
        self.assertIn("Раздел II", system)
        self.assertIn("существенн", system)
        self.assertIn("таблиц", system.lower())
        self.assertNotIn("Не вставляй таблицы", system)
        self.assertIn("Наблюдение 3.1", sections)
        self.assertIn("Наблюдения по итогам проверки", sections)
        self.assertIn("ровно 6", sections.lower())
        self.assertIn("Общая информация", sections)
        self.assertNotIn("сгруппируй", sections.lower())
        self.assertIn("{hypotheses_block}", prompt("conclusion_user"))
        self.assertIn("{observation_outline}", prompt("conclusion_user"))
        self.assertIn("подтверждённ", user)
        self.assertIn("Проверка аренды", user)
        self.assertIn("6 наблюдений", user)
        self.assertNotIn("{missing}", user)
        continue_user = prompt(
            "conclusion_continue_user",
            inspection="Проверка аренды",
            keywords="аренда",
            done_list="1, 3",
            hypothesis_count=4,
            hypothesis_numbers="4, 5, 6, 8",
            next_number="3.3",
            hypotheses_block="4. тест",
            program_block="",
            brief_block="",
            cards_block="",
            fragments="",
            general_tail=".",
        )
        self.assertIn("недостающ", continue_user.lower())
        self.assertIn("3.3", continue_user)


if __name__ == "__main__":
    unittest.main()
