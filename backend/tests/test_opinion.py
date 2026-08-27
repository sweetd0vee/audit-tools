import tempfile
import unittest
import zipfile
from pathlib import Path

from app.prompts import prompt
from app.services.brief_docx import write_opinion_docx
from app.services.hypotheses_flow import resolve_hypothesis_selection
from app.services.opinion_flow import (
    FONT_CALIBRI,
    FONT_TIMES,
    format_hypotheses_block,
    parse_document_font,
    parse_opinion_font_flag,
)


def _row(n: int, priority: str = "средний") -> dict:
    return {
        "n": str(n),
        "hypothesis": f"Гипотеза {n}",
        "priority": priority,
        "assertion": "полнота",
        "risk": "риск",
        "why_risk": "почему",
        "how_to_test": "сверка",
        "evidence_request": "договоры",
        "working_paper": "WP",
        "npa_criteria": "НК РБ",
        "plan_sections": "учёт",
    }


class TestOpinionFont(unittest.TestCase):
    def test_parse_flags(self):
        self.assertEqual(parse_document_font("c"), FONT_CALIBRI)
        self.assertEqual(parse_document_font("-c"), FONT_CALIBRI)
        self.assertEqual(parse_document_font("calibri"), FONT_CALIBRI)
        self.assertEqual(parse_document_font("t"), FONT_TIMES)
        self.assertEqual(parse_document_font("times"), FONT_TIMES)
        self.assertEqual(parse_document_font(None), FONT_TIMES)
        self.assertEqual(parse_opinion_font_flag("аудиторское мнение -c"), FONT_CALIBRI)
        self.assertEqual(parse_opinion_font_flag("аудиторское мнение -t заново"), FONT_TIMES)
        self.assertEqual(parse_opinion_font_flag("аудиторское мнение"), FONT_TIMES)
        self.assertEqual(parse_opinion_font_flag("мнение Calibri"), FONT_CALIBRI)


class TestHypothesisSelection(unittest.TestCase):
    def test_picks_numbers_and_high(self):
        rows = [_row(i, "высокий" if i <= 2 else "средний") for i in range(1, 9)]
        self.assertEqual(resolve_hypothesis_selection(rows, numbers=[1, 3, 5]), [1, 3, 5])
        self.assertEqual(resolve_hypothesis_selection(rows, all_high=True), [1, 2])
        self.assertEqual(len(resolve_hypothesis_selection(rows, all_rows=True)), 8)

    def test_rejects_unknown_number(self):
        rows = [_row(i) for i in range(1, 9)]
        with self.assertRaises(ValueError):
            resolve_hypothesis_selection(rows, numbers=[12])

    def test_rejects_empty(self):
        rows = [_row(i) for i in range(1, 9)]
        with self.assertRaises(ValueError):
            resolve_hypothesis_selection(rows, numbers=[])


class TestOpinionDocx(unittest.TestCase):
    def test_writes_title_font_and_skips_tables(self):
        body = """
## Название проверки
Проверка аренды коммерческой недвижимости.

## Цели и задачи аудита
Цели аудита:
- Оценка учёта аренды.
- Выработка рекомендаций.

## Инструменты проверки
Анализ ЛПА и сверка договоров с критериями НПА.

## Аудиторское мнение
По подтверждённым гипотезам отмечается риск неполного отражения платежей.

| таблица | не должна | попасть |
| --- | --- | --- |

## Рекомендации
- Сверить регистры аренды с договорами.
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "opinion.docx"
            write_opinion_docx(
                path,
                inspection_name="Проверка аренды коммерческой недвижимости",
                period="2025",
                keywords=["аренда"],
                case_id="abc123",
                body=body,
                font=FONT_CALIBRI,
            )
            self.assertTrue(path.exists())
            from docx import Document

            doc = Document(str(path))
            texts = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("I. Аудиторское мнение по итогам проверки", texts)
            self.assertIn("Проверка аренды коммерческой недвижимости", texts)
            self.assertIn("Анализ ЛПА", texts)
            self.assertIn("Сверить регистры", texts)
            self.assertNotIn("не должна", texts)
            self.assertEqual(len(doc.tables), 0)
            with zipfile.ZipFile(path) as zf:
                xml = zf.read("word/document.xml").decode("utf-8")
            self.assertIn("Calibri", xml)
            self.assertNotIn("w:tbl", xml)

    def test_keeps_case_title_and_drops_model_title_section(self):
        body = """
## Название проверки
Проверка соблюдения принципов налогообложения и требований нормативных правовых актов.

## Цели и задачи аудита
Цели аудита:
- Оценка учёта аренды.
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "opinion.docx"
            write_opinion_docx(
                path,
                inspection_name="Проверка аренды коммерческой недвижимости",
                period=None,
                keywords=["аренда"],
                case_id="c1",
                body=body,
                font=FONT_TIMES,
            )
            from docx import Document

            texts = "\n".join(p.text for p in Document(str(path)).paragraphs)
            self.assertIn("Проверка аренды коммерческой недвижимости", texts)
            self.assertNotIn("принципов налогообложения", texts)
            self.assertIn("Оценка учёта аренды", texts)

    def test_times_default_in_xml(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "opinion.docx"
            write_opinion_docx(
                path,
                inspection_name="Проверка кассы",
                period=None,
                keywords=[],
                case_id="c1",
                body="## Аудиторское мнение\nКраткий вывод по подтверждённым гипотезам.",
                font=FONT_TIMES,
            )
            with zipfile.ZipFile(path) as zf:
                xml = zf.read("word/document.xml").decode("utf-8")
            self.assertIn("Times New Roman", xml)


class TestOpinionPrompts(unittest.TestCase):
    def test_prompts_require_narrative_and_confirmed_only(self):
        system = prompt("opinion_system")
        sections = prompt("opinion_sections")
        user = prompt(
            "opinion_user",
            inspection="Проверка аренды",
            keywords="аренда",
            period="2025",
            document_catalog="",
            hypotheses_block="1. тест",
            program_block="",
            brief_block="",
            total_block="",
            cards_block="",
            fragments="",
            sections=sections,
            target=3600,
            target_hi=7200,
        )
        self.assertIn("подтверждённ", system)
        self.assertIn("дословн", system)
        self.assertIn("таблиц", system)
        self.assertIn("Цели и задачи аудита", sections)
        self.assertIn("Инструменты проверки", sections)
        self.assertIn("Аудиторское мнение", sections)
        self.assertIn("Рекомендации", sections)
        self.assertIn("{hypotheses_block}", prompt("opinion_user"))
        self.assertIn("подтверждённ", user)
        self.assertIn("дословн", user)

    def test_format_hypotheses_block(self):
        text = format_hypotheses_block([_row(1, "высокий"), _row(3, "средний")])
        self.assertIn("1. [высокий] Гипотеза 1", text)
        self.assertIn("3. [средний] Гипотеза 3", text)
        self.assertIn("Риск:", text)


if __name__ == "__main__":
    unittest.main()
