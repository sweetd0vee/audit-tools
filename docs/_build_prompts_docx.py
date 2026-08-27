"""Собрать Word со всеми промптами агента. Run: python docs/_build_prompts_docx.py"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
PROMPTS = ROOT / "prompts"
OUT = ROOT / "Промпты агента.docx"

NAVY = RGBColor(0x0B, 0x1F, 0x33)
INK = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x5A, 0x68, 0x76)
TEAL = RGBColor(0x1E, 0x5F, 0x58)
GOLD = RGBColor(0x8A, 0x6A, 0x18)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
PLACE = RGBColor(0x8A, 0x2E, 0x1A)

PLACE_RE = re.compile(r"(\{\{?[^{}]+\}?\})")

STEPS: list[dict] = [
    {
        "title": "Шаг 0. Справка агента",
        "command": "помощь / help / ?",
        "purpose": (
            "Это не промпт языковой модели. Это текст, который Pipe «Аудитор» "
            "показывает в чате, когда аудитор просит помощь или ещё ничего не написал. "
            "Здесь объясняется, как начать проверку и какие команды есть дальше."
        ),
        "result": "Короткий гид в чате: как описать проверку, как утвердить документы, что писать дальше.",
        "prompts": [
            {
                "file": "pipe_help.txt",
                "kind": "Текст Pipe (не LLM)",
                "label": "Справка в чате",
                "why": (
                    "Ответ на «помощь». Канон для правки — этот файл; после изменения "
                    "текст нужно скопировать в HELP в seed/openwebui/functions/audit_agent.py "
                    "и заново вставить Pipe в Open WebUI."
                ),
                "placeholders": [],
            }
        ],
    },
    {
        "title": "Шаг 1. Старт проверки — список НПА",
        "command": "Название проверки, ключевые слова (например: Проверка аренды…, аренда, валюта, НДС)",
        "purpose": (
            "Аудитор описал, что проверяет. Модель предлагает список нормативных актов, "
            "с которыми нужно ознакомиться до анализа данных. Ничего не скачивается, "
            "пока аудитор не напишет «утверждаю 1, 2, 4» (или «все обязательные»)."
        ),
        "result": "Нумерованный список актов в чате: название, приоритет, зачем нужен. Ответ модели — строго JSON.",
        "prompts": [
            {
                "file": "propose_system.txt",
                "kind": "System",
                "label": "Роль: методолог подбирает НПА",
                "why": (
                    "Задаёт роль и жёсткие рамки: только право РБ, релевантность теме, "
                    "поисковые запросы на официальных сайтах, без выдуманных номеров статей, "
                    "ответ только JSON."
                ),
                "placeholders": [],
            },
            {
                "file": "propose_user.txt",
                "kind": "User",
                "label": "Задание: собрать JSON со списком актов",
                "why": (
                    "Сюда подставляются название проверки, ключевые слова и период. "
                    "Здесь же схема JSON и диапазон числа документов. Не ломайте пример JSON "
                    "«красивым markdown» — модель должна вернуть именно JSON."
                ),
                "placeholders": [
                    "inspection_name",
                    "keywords_str",
                    "period_str",
                    "min_docs",
                    "max_docs",
                ],
            },
        ],
    },
    {
        "title": "Шаг 2. Обычный чат",
        "command": "Любой текст без служебной команды (не «саммари», не «вопрос …» и т.д.)",
        "purpose": (
            "Свободный диалог по проверке: план, формулировки, риски, черновики процедур. "
            "Модель не ходит в базу знаний кейса. Если нужна цитата из скачанного акта — "
            "аудитор должен начать сообщение со слова «вопрос»."
        ),
        "result": "Короткий ответ в чате, без Word-файла.",
        "prompts": [
            {
                "file": "chat_system.txt",
                "kind": "System",
                "label": "Роль: помощник в обычном диалоге",
                "why": (
                    "Системная роль для /api/v1/chat. User — история сообщений чата "
                    "(без меток кейса). Здесь можно усилить тон, запреты и когда "
                    "перенаправлять в «вопрос …»."
                ),
                "placeholders": [],
            }
        ],
    },
    {
        "title": "Шаг 3. Вопрос по базе знаний",
        "command": "вопрос …  /  вопрос по базе: …  /  /ask …",
        "purpose": (
            "Ответ строго по дословным фрагментам скачанных актов плюс карточки саммари "
            "как ориентир, какой это документ. Без префикса «вопрос» этот промпт не вызывается."
        ),
        "result": "Ответ в чате + блок «Откуда в базе знаний» с цитатами.",
        "prompts": [
            {
                "file": "ask_system.txt",
                "kind": "System",
                "label": "Роль: ассистент по фрагментам НПА",
                "why": (
                    "Главные запреты RAG: номер статьи и формулировку брать только из "
                    "дословных фрагментов; если ответа нет — так и сказать; не выдумывать нормы."
                ),
                "placeholders": [],
            },
            {
                "file": "ask_user.txt",
                "kind": "User",
                "label": "Задание: вопрос + фрагменты",
                "why": (
                    "Сюда кладутся тема проверки, вопрос аудитора, блок карточек саммари "
                    "и дословные фрагменты индекса. Обогащайте инструкцию, как цитировать "
                    "и когда отказываться."
                ),
                "placeholders": [
                    "inspection",
                    "keywords",
                    "question",
                    "summary_block",
                    "context",
                ],
            },
        ],
    },
    {
        "title": "Шаг 4. Саммари по базе знаний",
        "command": "саммари  /  сводка  /  /brief",
        "purpose": (
            "Карточки существенного по каждому скачанному акту и обзор нормативного контура "
            "в начале Word. Короткий акт читается целиком (oneshot). Длинный — через RAG-выборку "
            "фрагментов (rag_card). Затем один вызов overview по всем карточкам."
        ),
        "result": "Word: обзор контура + карточка на каждый акт (ссылки на статьи).",
        "prompts": [
            {
                "file": "summary_system.txt",
                "kind": "System",
                "label": "Роль на все вызовы карточки и обзора",
                "why": (
                    "Общая роль старшего аудитора: не конспект всех статей, а существенное "
                    "именно для этой проверки. Один system на oneshot, rag_card и overview."
                ),
                "placeholders": [],
            },
            {
                "file": "oneshot_card.txt",
                "kind": "User",
                "label": "Карточка короткого акта (весь текст)",
                "why": (
                    "Когда акт умещается в одно окно. Модель видит полный текст и оглавление. "
                    "Здесь задаётся структура карточки (## Зачем / Суть / Ключевые нормы / "
                    "Что проверять / Чего нет) и объём."
                ),
                "placeholders": ["inspection", "keywords", "title", "source", "outline", "body"],
            },
            {
                "file": "rag_card.txt",
                "kind": "User",
                "label": "Карточка длинного акта (RAG-выборка)",
                "why": (
                    "Когда акт длинный: в {body} не весь кодекс, а отобранные фрагменты [n]. "
                    "Модель не должна опираться на статьи, которых нет во фрагментах, "
                    "даже если они есть в оглавлении."
                ),
                "placeholders": [
                    "inspection",
                    "keywords",
                    "title",
                    "source",
                    "chars",
                    "outline",
                    "body",
                ],
            },
            {
                "file": "overview.txt",
                "kind": "User",
                "label": "Обзор нормативного контура (начало Word)",
                "why": (
                    "После карточек по всем актам. Сводит рамку проверки: как акты складываются, "
                    "на что опираться в первую очередь, где пересечения и пробелы. Не копирует карточки."
                ),
                "placeholders": ["inspection", "keywords", "cards"],
            },
        ],
    },
    {
        "title": "Шаг 5. Саммари total — из знаний модели",
        "command": "саммари total  /  саммари тотал  /  конспект модели  /  /total",
        "purpose": (
            "Конспект по теме проверки из знаний модели, без файлов кейса и без базы знаний. "
            "Это «что модель помнит про право РБ», чтобы сверить с саммари по актам. "
            "Номера статей модель может помнить неточно — промпт требует пометки «уточнить»."
        ),
        "result": "Word-конспект с разделами и списком источников [n].",
        "prompts": [
            {
                "file": "total_system.txt",
                "kind": "System",
                "label": "Роль: конспект из знаний модели",
                "why": (
                    "Запрещает опираться на файлы кейса, требует право РБ, ссылки [n], "
                    "раздел ## Источники и отказ от суждения «нарушение / не нарушение»."
                ),
                "placeholders": [],
            },
            {
                "file": "total_sections.txt",
                "kind": "Шаблон разделов",
                "label": "Какие разделы должны быть в Word",
                "why": (
                    "Вставляется внутрь user-промпта как {sections}. Здесь правите состав "
                    "и наполнение глав конспекта (суть, акты, важное, акценты, ограничения, источники)."
                ),
                "placeholders": [],
            },
            {
                "file": "total_user.txt",
                "kind": "User",
                "label": "Задание: собрать конспект по теме",
                "why": (
                    "Обёртка: название проверки, ключевые слова, период, вставка разделов, "
                    "целевой объём. Явно сказано не использовать тексты файлов кейса."
                ),
                "placeholders": [
                    "inspection",
                    "keywords",
                    "period",
                    "sections",
                    "target",
                    "target_hi",
                ],
            },
        ],
    },
    {
        "title": "Шаг 6. Программа проверки",
        "command": "программа проверки  /  программа проверки 8  /  программа проверки 10-12",
        "purpose": (
            "Черновик программы СВА: таблица «вопросы, подлежащие аудиту», как в типовой "
            "программе банка. Не акт и не заключение. Число пунктов задаёт аудитор; "
            "без числа — обычно 8–11. Каждый пункт — 3–4 предложения: что проверить, "
            "как сверить, какие документы, критерий [n]."
        ),
        "result": "Word-таблица вопросов программы (черновик, не подпись руководителя СВА).",
        "prompts": [
            {
                "file": "program_system.txt",
                "kind": "System",
                "label": "Роль: аудитор пишет программу СВА",
                "why": (
                    "Рамки рабочей программы: право РБ, без суждения, стиль номинативных "
                    "формулировок, жёсткий лимит числа пунктов, типовая логика направлений проверки."
                ),
                "placeholders": [],
            },
            {
                "file": "program_sections.txt",
                "kind": "Шаблон разделов",
                "label": "Формат ответа: название, период, вопросы",
                "why": (
                    "Канон markdown, из которого сервер собирает Word-таблицу. Здесь правите "
                    "как звучит пункт (3–4 предложения, критерий, документы). {items_hint} "
                    "подставляет «строго N» или диапазон."
                ),
                "placeholders": ["items_hint"],
            },
            {
                "file": "program_user.txt",
                "kind": "User",
                "label": "Задание: собрать программу по материалам кейса",
                "why": (
                    "Сюда идут каталог документов, фрагменты НПА, карточки саммари и лимит пунктов. "
                    "Обогащайте, если хотите жёстче привязать вопросы к актам или к теме проверки."
                ),
                "placeholders": [
                    "inspection",
                    "keywords",
                    "period",
                    "document_catalog",
                    "catalog",
                    "fragments",
                    "cards_block",
                    "sections",
                    "items_hint",
                    "target",
                    "target_hi",
                ],
            },
        ],
    },
    {
        "title": "Шаг 7. Гипотезы",
        "command": "гипотезы  /  чеклист  /  /hypotheses",
        "purpose": (
            "Чеклист 8–10 проверяемых гипотез для планирования. Контекст: карточки саммари, "
            "total, программа, фрагменты НПА и знания модели. Это не выводы проверки, "
            "а предположения о риске / слабости контроля, которые аудитор потом подтвердит."
        ),
        "result": "Excel-чеклист. Дальше: «утверждаю гипотезы 1, 3, 5» — в мнение пойдут только они.",
        "prompts": [
            {
                "file": "hypotheses_system.txt",
                "kind": "System",
                "label": "Роль: формулировать проверяемые гипотезы",
                "why": (
                    "Качество гипотезы: одно предложение «если …, то риск …»; конкретный риск "
                    "для банка; метод проверки; что запросить; assertion аудита; без суждения "
                    "«нарушение установлено»."
                ),
                "placeholders": [],
            },
            {
                "file": "hypotheses_schema.txt",
                "kind": "Схема JSON",
                "label": "Формат ответа: поля чеклиста",
                "why": (
                    "Сервер парсит строго этот JSON. Не меняйте имена полей "
                    "(hypothesis, assertion, risk, …) без правки кода. Можно уточнять, "
                    "как заполнять каждое поле."
                ),
                "placeholders": [],
            },
            {
                "file": "hypotheses_user.txt",
                "kind": "User",
                "label": "Задание: 8–10 гипотез по материалам кейса",
                "why": (
                    "Сводка всех материалов + вставка схемы. Требование «только JSON» "
                    "критично: markdown вокруг сломает разбор."
                ),
                "placeholders": [
                    "inspection",
                    "keywords",
                    "period",
                    "document_catalog",
                    "catalog",
                    "fragments",
                    "cards_block",
                    "total_block",
                    "program_block",
                    "schema",
                ],
            },
        ],
    },
    {
        "title": "Шаг 8. Аудиторское мнение (раздел I)",
        "command": "аудиторское мнение  /  аудиторское мнение -c  /  -t",
        "purpose": (
            "Черновик раздела I заключения: «I. Аудиторское мнение по итогам проверки». "
            "2–4 страницы повествования для руководства банка. В текст попадают только "
            "подтверждённые гипотезы. Шрифт: -c Calibri, -t Times New Roman (по умолчанию)."
        ),
        "result": "Word раздела I. Нужен шаг «утверждаю гипотезы …».",
        "prompts": [
            {
                "file": "opinion_system.txt",
                "kind": "System",
                "label": "Роль: черновик раздела I",
                "why": (
                    "Агрегированный взгляд, без таблиц, без выдуманных фактов проверки, "
                    "без подписи. Название проверки копировать дословно. Только подтверждённые гипотезы."
                ),
                "placeholders": [],
            },
            {
                "file": "opinion_sections.txt",
                "kind": "Шаблон разделов",
                "label": "Структура: цели, инструменты, мнение, рекомендации",
                "why": (
                    "Канон глав раздела I. Здесь правите, как звучит мнение для руководства "
                    "и как сворачивать гипотезы в повествование."
                ),
                "placeholders": [],
            },
            {
                "file": "opinion_user.txt",
                "kind": "User",
                "label": "Задание: написать раздел I по подтверждённым гипотезам",
                "why": (
                    "Сюда кладутся гипотезы, программа, саммари, total, фрагменты НПА. "
                    "Жёсткий ориентир объёма 2–4 страницы."
                ),
                "placeholders": [
                    "inspection",
                    "keywords",
                    "period",
                    "document_catalog",
                    "hypotheses_block",
                    "program_block",
                    "brief_block",
                    "total_block",
                    "cards_block",
                    "fragments",
                    "sections",
                    "target",
                    "target_hi",
                ],
            },
        ],
    },
    {
        "title": "Шаг 9. Аудиторское заключение",
        "command": "аудиторское заключение  /  аудиторское заключение -c  /  -t",
        "purpose": (
            "Черновик полного заключения. Раздел I уже готов — модель его не пишет заново. "
            "Раздел II (таблица результатов) сервер соберёт сам — модель его тоже не пишет. "
            "Модель пишет наблюдения раздела III по каждой подтверждённой гипотезе и раздел IV "
            "«Общая информация». Шрифт тот же: -c / -t."
        ),
        "result": "Word полного черновика АЗ. Нужны гипотезы, их утверждение и аудиторское мнение.",
        "prompts": [
            {
                "file": "conclusion_system.txt",
                "kind": "System",
                "label": "Роль: наблюдения III + общая информация IV",
                "why": (
                    "Каждая подтверждённая гипотеза = отдельное наблюдение с уровнем "
                    "существенности и рекомендацией. Не выдумывать реквизиты приказа, "
                    "состав группы, сроки — «уточняется / заполняет аудитор»."
                ),
                "placeholders": [],
            },
            {
                "file": "conclusion_sections.txt",
                "kind": "Шаблон разделов",
                "label": "Канон наблюдений 3.1, 3.2… и полей раздела IV",
                "why": (
                    "Жёсткий markdown, из которого собирается Word. Меняя заголовки ### Наблюдение, "
                    "существенность / гипотеза / рекомендация — сверяйтесь с парсером в коде."
                ),
                "placeholders": [],
            },
            {
                "file": "conclusion_user.txt",
                "kind": "User",
                "label": "Задание: III+IV по гипотезам и готовому мнению",
                "why": (
                    "Готовый раздел I передаётся для согласованности формулировок, копировать "
                    "его в ответ нельзя. Писать только III и IV."
                ),
                "placeholders": [
                    "inspection",
                    "keywords",
                    "period",
                    "document_catalog",
                    "hypotheses_block",
                    "opinion_block",
                    "program_block",
                    "brief_block",
                    "total_block",
                    "cards_block",
                    "fragments",
                    "sections",
                ],
            },
        ],
    },
]

APPENDIX: list[dict] = [
    {
        "title": "Приложение А. Старый map-reduce саммари (сейчас не вызывается)",
        "purpose": (
            "Раньше длинный акт резали на части: map_essential по куску, retry если пусто, "
            "reduce_card склеивал заметки. Сейчас длинный акт идёт через RAG (rag_card). "
            "Файлы лежат в docs/prompts/ — можно взять идеи или вернуть контур."
        ),
        "prompts": [
            {
                "file": "map_essential.txt",
                "kind": "User (не используется)",
                "label": "Выжимка существенного из одной части акта",
                "why": "Map-шаг: из куска текста — только нормы, нужные этой проверке.",
                "placeholders": ["inspection", "keywords", "idx", "total", "title", "outline", "body"],
            },
            {
                "file": "retry_essential.txt",
                "kind": "User (не используется)",
                "label": "Повтор, если заметки части пустые",
                "why": "Второй заход, когда модель ничего существенного не нашла.",
                "placeholders": [
                    "inspection",
                    "keywords",
                    "idx",
                    "total",
                    "title",
                    "previous",
                    "body",
                ],
            },
            {
                "file": "reduce_card.txt",
                "kind": "User (не используется)",
                "label": "Сборка карточки из заметок по всем частям",
                "why": "Reduce-шаг: одна карточка из map-заметок, без склейки «как есть».",
                "placeholders": [
                    "inspection",
                    "keywords",
                    "title",
                    "chars",
                    "parts",
                    "notes",
                ],
            },
        ],
    },
    {
        "title": "Приложение Б. Open WebUI, не backend",
        "purpose": (
            "Эти тексты не читает FastAPI при саммари/программе. Они для настройки чата "
            "Open WebUI (RAG-шаблон) и запасного контура Ollama + Tools без Pipe."
        ),
        "prompts": [
            {
                "file": "rag_template.txt",
                "kind": "Шаблон RAG Open WebUI",
                "label": "Как модель должна использовать {{CONTEXT}}",
                "why": (
                    "Срабатывает, если в чат подмешивают Knowledge Open WebUI. "
                    "{{CONTEXT}} — служебное, не удалять. Дубли: seed/openwebui/RAG_TEMPLATE.txt "
                    "и docker-compose.yml (RAG_TEMPLATE). На живом volume ещё Admin → Documents."
                ),
                "placeholders": [],
            },
            {
                "file": "system_auditor.txt",
                "kind": "System Tools (не Pipe)",
                "label": "Роль, если tools цепляют к голой Ollama",
                "why": (
                    "Продуктовый вход — Pipe. Этот файл нужен только для Ollama + Tools "
                    "(audit_case.py). Копия: seed/openwebui/SYSTEM_AUDITOR.txt."
                ),
                "placeholders": [],
            },
        ],
    },
    {
        "title": "Приложение В. Дописывание заключения (код пока не вызывает)",
        "purpose": (
            "Заготовка user-промпта, если модель оборвала раздел III и не покрыла все "
            "подтверждённые гипотезы. Файл лежит в docs/prompts/, но backend его сейчас "
            "не читает. Можно обогатить и потом включить в conclusion_flow."
        ),
        "prompts": [
            {
                "file": "conclusion_continue_user.txt",
                "kind": "User (не используется)",
                "label": "Дописать недостающие наблюдения",
                "why": (
                    "Продолжение после обрыва: уже написанные гипотезы не повторять, "
                    "нумерация с ### Наблюдение {next_number}."
                ),
                "placeholders": [
                    "inspection",
                    "keywords",
                    "done_list",
                    "hypothesis_count",
                    "hypothesis_numbers",
                    "next_number",
                    "hypotheses_block",
                    "program_block",
                    "brief_block",
                    "cards_block",
                    "fragments",
                    "general_tail",
                ],
            },
        ],
    },
]


def load_prompt(name: str) -> str:
    path = PROMPTS / name
    return path.read_text(encoding="utf-8").replace("\r\n", "\n").strip("\n")


def set_run_font(run, *, size=12, bold=False, italic=False, color=INK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), font)
    r_fonts.set(qn("w:cs"), font)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_margins(cell, *, top=80, bottom=80, left=120, right=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def set_cell_borders(cell, color: str = "C8D0D8", sz: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def no_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def keep_together(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def style_heading(paragraph, level: int) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(18 if level == 1 else 12)
    fmt.space_after = Pt(8)
    fmt.line_spacing = 1.15
    if level == 1:
        fmt.page_break_before = True


def add_text(paragraph, text, **kwargs) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, **kwargs)


def add_body(doc: Document, text: str, *, italic=False, size=12, space_after=8) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    add_text(p, text, size=size, italic=italic)


def add_label_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    add_text(p, label, size=11, bold=True, color=TEAL)
    add_text(p, value, size=11, color=INK)


def boxed(doc: Document, text: str, *, fill: str, border: str, italic=False, size=11) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    shade_cell(cell, fill)
    set_cell_borders(cell, border, "12")
    cell_margins(cell, top=100, bottom=100, left=140, right=140)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    add_text(p, text, size=size, italic=italic, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def write_prompt_lines(cell, text: str) -> None:
    cell.text = ""
    lines = text.split("\n")
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        if not line:
            add_text(p, " ", size=10, font="Consolas", color=INK)
            continue
        pos = 0
        for match in PLACE_RE.finditer(line):
            if match.start() > pos:
                add_text(p, line[pos : match.start()], size=10, font="Consolas", color=INK)
            add_text(
                p,
                match.group(1),
                size=10,
                bold=True,
                font="Consolas",
                color=PLACE,
            )
            pos = match.end()
        if pos < len(line):
            add_text(p, line[pos:], size=10, font="Consolas", color=INK)


def add_prompt_box(doc: Document, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    shade_cell(cell, "F6F3EC")
    set_cell_borders(cell, "C4B89A", "12")
    cell_margins(cell, top=120, bottom=120, left=140, right=140)
    write_prompt_lines(cell, body)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_note_box(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F7F6")
    set_cell_borders(cell, "B7D0C8", "8")
    cell_margins(cell, top=80, bottom=140, left=140, right=140)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_text(p, "Комментарий / идея правки", size=10, bold=True, italic=True, color=TEAL)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(
        p2,
        "Пишите сюда, что усилить: роль, запреты, примеры хорошего ответа, объём, тон. "
        "Плейсхолдеры {в фигурных скобках} не переименовывайте.",
        size=10,
        italic=True,
        color=MUTED,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(10)


def add_heading_styled(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    style_heading(p, level)
    for run in p.runs:
        set_run_font(
            run,
            size=18 if level == 1 else 14,
            bold=True,
            color=NAVY if level == 1 else TEAL,
            font="Calibri",
        )


def add_cover(doc: Document) -> None:
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    add_text(p, "Агент «Аудитор»", size=14, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    add_text(p, "Промпты по шагам проверки", size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    add_text(
        p,
        "Рабочий документ для правки и обогащения.\n"
        "Полные тексты в том порядке, в котором идёт агент в чате.",
        size=13,
        italic=True,
        color=MUTED,
    )

    boxed(
        doc,
        "Как пользоваться. Идите по шагам сверху вниз — так аудитор ведёт проверку. "
        "У каждого шага: зачем он, что писать в чате, какой файл получается. "
        "Дальше — сами промпты: System (роль модели), User (задание с подстановками), "
        "шаблон разделов или JSON. Красным выделены плейсхолдеры {имя}: код подставляет "
        "их как есть, переименовывать нельзя. Под каждым промптом — поле для ваших правок. "
        "Чтобы агент начал использовать новый текст, перенесите его в docs/prompts/<файл>.txt "
        "(сервер читает эти файлы при каждом вызове).",
        fill="EAF3F1",
        border="2F8F84",
        size=11,
    )


def add_howto(doc: Document) -> None:
    add_heading_styled(doc, "Как устроены промпты", 1)
    add_body(
        doc,
        "На одном шаге обычно несколько текстов. Их видит модель вместе: сначала system, "
        "потом user. Шаблон разделов или JSON-схема вставляется внутрь user как {sections} / {schema}.",
    )
    rows = [
        ("System", "Роль и запреты на весь шаг. Меняйте тон, право РБ, что нельзя выдумывать."),
        ("User", "Конкретное задание. Сюда код подставляет тему проверки, фрагменты, карточки."),
        ("Шаблон разделов", "Какие ## заголовки ждать в ответе. От них зависит сборка Word."),
        ("Схема JSON", "Поля Excel-чеклиста гипотез и списка НПА. Имена полей трогать опасно."),
        ("Текст Pipe", "Не модель: реплика агента в чате (справка)."),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.autofit = True
    hdr = table.rows[0].cells
    for i, title in enumerate(("Тип", "Зачем править")):
        hdr[i].text = ""
        shade_cell(hdr[i], "0B1F33")
        cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        add_text(p, title, size=11, bold=True, color=WHITE)
    for r, (kind, why) in enumerate(rows, start=1):
        fill = "FFFFFF" if r % 2 else "F4F7F9"
        c0, c1 = table.rows[r].cells
        for cell, txt, bold in ((c0, kind, True), (c1, why, False)):
            cell.text = ""
            shade_cell(cell, fill)
            cell_margins(cell)
            set_cell_borders(cell, "D5DEE6", "4")
            add_text(cell.paragraphs[0], txt, size=11, bold=bold, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_body(
        doc,
        "Цепочка в чате (по порядку): описать проверку → утвердить документы → "
        "программа / саммари / саммари total (можно в любом порядке) → гипотезы → "
        "утвердить гипотезы → аудиторское мнение → аудиторское заключение. "
        "«вопрос …» и обычный чат доступны, когда кейс уже есть.",
        italic=True,
    )


def add_map(doc: Document) -> None:
    add_heading_styled(doc, "Карта шагов", 1)
    add_body(doc, "Кратко: команда аудитора, что получается, какие файлы промптов.")
    rows = [
        ("0", "помощь", "справка в чате", "pipe_help"),
        ("1", "название проверки", "список НПА", "propose_system + propose_user"),
        ("2", "обычный текст", "диалог", "chat_system"),
        ("3", "вопрос …", "ответ с цитатой", "ask_system + ask_user"),
        ("4", "саммари", "Word по базе", "summary_system + oneshot / rag_card + overview"),
        ("5", "саммари total", "Word из знаний модели", "total_system + total_user + total_sections"),
        ("6", "программа проверки", "Word-таблица вопросов", "program_system + user + sections"),
        ("7", "гипотезы", "Excel-чеклист", "hypotheses_system + user + schema"),
        ("8", "аудиторское мнение", "Word, раздел I", "opinion_system + user + sections"),
        ("9", "аудиторское заключение", "Word, полный черновик", "conclusion_system + user + sections"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    headers = ("Шаг", "В чате", "Результат", "Промпты")
    for i, title in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        shade_cell(cell, "0B1F33")
        cell_margins(cell, top=60, bottom=60, left=80, right=80)
        add_text(cell.paragraphs[0], title, size=10, bold=True, color=WHITE)
    for r, row in enumerate(rows, start=1):
        fill = "FFFFFF" if r % 2 else "F4F7F9"
        for i, txt in enumerate(row):
            cell = table.rows[r].cells[i]
            cell.text = ""
            shade_cell(cell, fill)
            cell_margins(cell, top=50, bottom=50, left=80, right=80)
            set_cell_borders(cell, "D5DEE6", "4")
            add_text(cell.paragraphs[0], txt, size=10, bold=(i == 0), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_prompt_block(doc: Document, spec: dict) -> None:
    add_heading_styled(doc, f"{spec['label']}  ·  {spec['file']}", 2)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(6)
    add_text(meta, spec["kind"], size=10, bold=True, color=GOLD)
    add_text(meta, "   файл  docs/prompts/", size=10, color=MUTED)
    add_text(meta, spec["file"], size=10, bold=True, color=MUTED)
    if spec.get("placeholders"):
        add_text(meta, "   плейсхолдеры: ", size=10, color=MUTED)
        add_text(
            meta,
            ", ".join("{" + p + "}" for p in spec["placeholders"]),
            size=10,
            bold=True,
            color=PLACE,
        )
    add_body(doc, spec["why"], size=11)
    add_prompt_box(doc, load_prompt(spec["file"]))
    add_note_box(doc)


def add_step(doc: Document, step: dict) -> None:
    add_heading_styled(doc, step["title"], 1)
    add_label_line(doc, "Команда в чате.  ", step["command"])
    add_label_line(doc, "Что получается.  ", step["result"])
    boxed(doc, "Для чего.  " + step["purpose"], fill="EAF3F1", border="2F8F84", size=11)
    for spec in step["prompts"]:
        add_prompt_block(doc, spec)


def set_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    r_pr = normal.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Calibri")
    r_fonts.set(qn("w:cs"), "Calibri")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Промпты агента «Аудитор»  ·  ", size=9, color=MUTED)
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    set_run_font(run, size=9, color=MUTED)

    core = doc.core_properties
    core.title = "Промпты агента «Аудитор»"
    core.author = "audit-tools"
    core.subject = "Все промпты по шагам проверки, для правки и обогащения"


def build() -> Path:
    missing = []
    for group in STEPS + APPENDIX:
        for spec in group["prompts"]:
            if not (PROMPTS / spec["file"]).is_file():
                missing.append(spec["file"])
    if missing:
        raise FileNotFoundError("Нет файлов промптов: " + ", ".join(missing))

    extra = {p.name for p in PROMPTS.glob("*.txt")}
    known = {spec["file"] for g in STEPS + APPENDIX for spec in g["prompts"]}
    leftover = sorted(extra - known)
    if leftover:
        raise RuntimeError("Промпты не попали в документ: " + ", ".join(leftover))

    doc = Document()
    set_styles(doc)
    add_cover(doc)
    add_howto(doc)
    add_map(doc)
    for step in STEPS:
        add_step(doc, step)
    for block in APPENDIX:
        add_heading_styled(doc, block["title"], 1)
        boxed(doc, block["purpose"], fill="F7F1E4", border="C4962C", size=11)
        for spec in block["prompts"]:
            add_prompt_block(doc, spec)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(path)
