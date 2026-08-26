from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

p = Path(r"c:\Users\audit\Work\Arina\2026\audit-tools\docs\программа.docx")
d = Document(str(p))
out = Path(r"c:\Users\audit\Work\Arina\2026\audit-tools\docs\_program_extract.txt")
lines: list[str] = []


def dump_para(para: Paragraph, indent: str = "") -> None:
    style = para.style.name if para.style else ""
    pf = para.paragraph_format
    lines.append(
        f"{indent}P style={style!r} align={para.alignment} "
        f"sb={pf.space_before} sa={pf.space_after} ls={pf.line_spacing} "
        f"li={pf.left_indent} text={para.text!r}"
    )
    for run in para.runs:
        f = run.font
        color = None
        try:
            color = str(f.color.rgb) if f.color and f.color.rgb else None
        except Exception:
            color = None
        lines.append(
            f"{indent}  RUN bold={run.bold} italic={run.italic} underline={run.underline} "
            f"size={f.size} name={f.name} color={color} text={run.text!r}"
        )


def dump_cell(cell, indent: str = "") -> None:
    tc = cell._tc
    tcPr = tc.tcPr
    if tcPr is not None:
        shd = tcPr.find(qn("w:shd"))
        if shd is not None:
            lines.append(f"{indent}SHD fill={shd.get(qn('w:fill'))} val={shd.get(qn('w:val'))}")
        gridSpan = tcPr.find(qn("w:gridSpan"))
        if gridSpan is not None:
            lines.append(f"{indent}GRIDSPAN={gridSpan.get(qn('w:val'))}")
        vMerge = tcPr.find(qn("w:vMerge"))
        if vMerge is not None:
            lines.append(f"{indent}VMERGE={vMerge.get(qn('w:val'))}")
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is not None:
            lines.append(f"{indent}WIDTH={tcW.get(qn('w:w'))} type={tcW.get(qn('w:type'))}")
        vAlign = tcPr.find(qn("w:vAlign"))
        if vAlign is not None:
            lines.append(f"{indent}VALIGN={vAlign.get(qn('w:val'))}")
    for para in cell.paragraphs:
        dump_para(para, indent)


def dump_table(table: Table, indent: str = "") -> None:
    lines.append(
        f"{indent}TABLE rows={len(table.rows)} cols={len(table.columns)} "
        f"style={table.style.name if table.style else None}"
    )
    tblPr = table._tbl.tblPr
    if tblPr is not None:
        jc = tblPr.find(qn("w:jc"))
        if jc is not None:
            lines.append(f"{indent}  jc={jc.get(qn('w:val'))}")
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is not None:
            lines.append(f"{indent}  tblW={tblW.get(qn('w:w'))} type={tblW.get(qn('w:type'))}")
        look = tblPr.find(qn("w:tblLook"))
        if look is not None:
            lines.append(f"{indent}  tblLook={dict(look.attrib)}")
        borders = tblPr.find(qn("w:tblBorders"))
        if borders is not None:
            for child in borders:
                tag = child.tag.split("}")[-1]
                lines.append(f"{indent}  border {tag}={dict(child.attrib)}")
    grid = table._tbl.tblGrid
    if grid is not None:
        widths = [gc.get(qn("w:w")) for gc in grid]
        lines.append(f"{indent}  grid={widths}")
    seen = set()
    for ri, row in enumerate(table.rows):
        trPr = row._tr.trPr
        extra = ""
        if trPr is not None:
            trH = trPr.find(qn("w:trHeight"))
            if trH is not None:
                extra += f" h={trH.get(qn('w:val'))}"
        lines.append(f"{indent}  ROW {ri}{extra}")
        for ci, cell in enumerate(row.cells):
            cid = id(cell._tc)
            if cid in seen:
                lines.append(f"{indent}    C{ci}: (merged/repeat)")
                continue
            seen.add(cid)
            lines.append(f"{indent}    C{ci}:")
            dump_cell(cell, indent + "      ")


lines.append("=== BODY ORDER ===")
body = d.element.body
for i, child in enumerate(body):
    tag = child.tag.split("}")[-1]
    if tag == "p":
        para = Paragraph(child, d)
        dump_para(para)
    elif tag == "tbl":
        dump_table(Table(child, d))
    else:
        lines.append(f"{tag}")

lines.append("")
lines.append("=== SECTIONS ===")
for s in d.sections:
    lines.append(
        f"page {s.page_width.cm:.2f}x{s.page_height.cm:.2f} cm, "
        f"margins L{s.left_margin.cm:.2f} R{s.right_margin.cm:.2f} "
        f"T{s.top_margin.cm:.2f} B{s.bottom_margin.cm:.2f}"
    )
    if s.header.paragraphs:
        for hp in s.header.paragraphs:
            lines.append(f"HEADER {hp.text!r}")
    if s.footer.paragraphs:
        for fp in s.footer.paragraphs:
            lines.append(f"FOOTER {fp.text!r}")

out.write_text("\n".join(lines), encoding="utf-8")
print("wrote", out, "size", out.stat().st_size)
