# -*- coding: utf-8 -*-
"""Regenerate docs/План_локальный_Cursor_аудита_банка.docx"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DOCS = Path(__file__).resolve().parent
OUT = DOCS / "План_локальный_Cursor_аудита_банка.docx"
OUT_FALLBACK = DOCS / "План_локальный_Cursor_аудита_банка_v2.docx"

NAVY = RGBColor(0x1B, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x5A, 0x88)
MUTED = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x22, 0x22, 0x22)
HEADER_BG = "1B3A5F"
HEADER_FG = "FFFFFF"
ALT_ROW = "F4F7FA"


def set_run_font(run, name="Calibri", size=11, bold=False, color=BLACK, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def p_style(p, space_after=8, space_before=0, line=1.15):
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 2 else ACCENT
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(12)
    return h


def add_para(doc, text, *, bold=False, italic=False, size=11, space_after=8, color=BLACK):
    p = doc.add_paragraph()
    p_style(p, space_after=space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_rich(doc, parts, *, space_after=8):
    """parts: list of (text, kwargs)"""
    p = doc.add_paragraph()
    p_style(p, space_after=space_after)
    for text, kw in parts:
        run = p.add_run(text)
        set_run_font(run, **kw)
    return p


def add_bullet(doc, text, *, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.25 + 0.6 * level)
    p_style(p, space_after=4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text)
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p_style(p, space_after=4)
    r = p.add_run(text)
    set_run_font(r)
    return p


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=BLACK, size=10, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p_style(p, space_after=2, space_before=2, line=1.08)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        shade_cell(cell, HEADER_BG)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            set_cell_text(cell, val, size=10)
            if r_i % 2 == 1:
                shade_cell(cell, ALT_ROW)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_callout(doc, text):
    p = doc.add_paragraph()
    p_style(p, space_after=10, space_before=4)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    set_run_font(run, italic=True, size=11, color=NAVY)
    return p


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:eastAsia"), "Calibri")

    # --- Title ---
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_style(t, space_after=4)
    r = t.add_run("Локальный Cursor для внутреннего аудита банка (РБ)")
    set_run_font(r, size=22, bold=True, color=NAVY)

    st = doc.add_paragraph()
    p_style(st, space_after=4)
    r = st.add_run(
        "План реализации и качества ответов. Сборка: Open WebUI + Ollama + SearXNG + Audit Tool Server."
    )
    set_run_font(r, size=12, color=ACCENT)

    meta = doc.add_paragraph()
    p_style(meta, space_after=14)
    r = meta.add_run(
        "Версия 2.0  ·  24.08.2026  ·  заменяет план 1.0 от 20.08.2026  ·  "
        "клиентские данные и формулировки проверки — только on-prem"
    )
    set_run_font(r, size=9, italic=True, color=MUTED)

    add_callout(
        doc,
        "Главная цель не меняется: инструмент, которым аудитор банка ведёт проверку "
        "так же уверенно, как разработчик пишет код в Cursor. Не «ещё один чат с LLM». "
        "Планка — качество ответа, grounding в утверждённых актах, HITL и коробка, "
        "которая просто работает.",
    )

    # ========== 0 ==========
    add_heading_styled(doc, "0. Цель, инвариант, планка Cursor", 1)
    add_para(
        doc,
        "Внутренний аудитор банка каждый раз заново собирает НПА, читает кодексы, "
        "проектирует тесты и пишет working papers. Модель без документов галлюцинирует "
        "статьи. Документы без модели не масштабируются. Продукт закрывает разрыв: "
        "модель думает, утверждённые акты доказывают, аудитор подтверждает.",
    )
    add_rich(
        doc,
        [
            (
                "Инвариант (не ломаем). ",
                {"bold": True},
            ),
            (
                "Модель не заменяет чтение первоисточника и не ставит аудиторское суждение. "
                "Она ускоряет поиск нормы, черновик теста и черновик WP. Цитата проверяется "
                "по файлу в кейсе. Клиентские данные в интернет не уходят.",
                {},
            ),
        ],
    )

    add_heading_styled(doc, "0.1 Что значит «как Cursor» — это не метафора про UI", 2)
    add_para(
        doc,
        "Cursor победил не тем, что написал IDE с нуля. Он взял VS Code и сделал "
        "законченный продукт: мнение из коробки, grounding в этом репозитории, "
        "агент с руками, предсказуемый цикл «план → tool → проверка». Мы делаем "
        "то же со стеком Open WebUI + Ollama + SearXNG + тонкий Audit Tool Server.",
    )
    add_table(
        doc,
        ["Cursor", "Мы (целевое состояние)"],
        [
            [
                "Поставил — открыл — пишешь код",
                "Поставил — открыл — ведёшь проверку в одном чате",
            ],
            [
                "VS Code не форкали, обвязали",
                "Open WebUI не форкаем, засеваем (промпт, RAG, Pipe)",
            ],
            [
                "Цитата из этого репозитория, не «вообще из интернета»",
                "Цитата из этого утверждённого акта, не «из памяти модели»",
            ],
            [
                "grep + семантический поиск + чтение файла",
                "hybrid retrieval по статьям + открытие первоисточника",
            ],
            [
                "Терминал / линтер считают, модель не «примерно компилирует»",
                "SQL/rules считают факт; модель не «примерно вспоминает курс»",
            ],
            [
                "Один вход: окно редактора",
                "Один вход: чат Open WebUI + файлы кейса на диске",
            ],
            [
                "Не заставляет собирать стек руками",
                "compose up / инсталлятор, не вики из шести сервисов",
            ],
            [
                "Правила проекта (.cursor, AGENTS.md)",
                "Пакет проверки: keywords, known_sources, золотые вопросы, шаблон WP",
            ],
        ],
    )
    add_para(
        doc,
        "Если аудитор должен прочитать README, выбрать embedding, включить hybrid search "
        "и привязать Knowledge — это стенд разработчика, не продукт.",
        italic=True,
    )

    add_heading_styled(doc, "0.2 Что должно получиться (горизонт пилота)", 2)
    add_bullet(doc, "Аудитор открывает один URL, выбирает модель «Аудитор», описывает проверку.")
    add_bullet(
        doc,
        "Система предлагает список НПА. Ничего не качается, пока аудитор не напишет «утверждаю 1, 2, 4».",
    )
    add_bullet(
        doc,
        "Официальные тексты лежат в папке кейса. Саммари Word на 6–10 страниц со ссылками на статьи.",
    )
    add_bullet(
        doc,
        "Вопрос по норме → ответ с цитатой из утверждённой библиотеки. Нет фрагмента — отказ, не выдуманная статья.",
    )
    add_bullet(
        doc,
        "Дальше (ещё не готово): факт из Excel через DuckDB, черновики тестов и WP, детерминированные rules, audit trail.",
    )
    add_bullet(
        doc,
        "Контур банка: Ollama / Open WebUI / SearXNG / данные кейса — on-prem. SearXNG только allowlist доменов РБ.",
    )

    # ========== 1 ==========
    add_heading_styled(doc, "1. Реализация: что уже собрано (факт на 24.08.2026)", 1)
    add_para(
        doc,
        "План 1.0 от 20.08 предлагал начать с Excel-tools (data_profile) и 10 автономных фаз. "
        "Это было бы «умный чат по таблицам» без нормативного контура. Мы сознательно "
        "перевернули порядок: сначала библиотека НПА и цитата — это аналог «модель видит "
        "этот репозиторий» у Cursor. Без этого любой finding юридически пустой.",
    )

    add_heading_styled(doc, "1.1 Стек в репозитории audit-tools", 2)
    add_table(
        doc,
        ["Компонент", "Роль", "Свой код", "Статус"],
        [
            [
                "Open WebUI :3000",
                "Единственная поверхность аудитора: чат, Knowledge, auth, Files",
                "Нет — ставим и засеваем",
                "есть",
            ],
            [
                "Pipe «Аудитор»",
                "Фазы и HITL в Python. 35B плохо держит native tool-calling — фазу выбирает код",
                "seed/openwebui/functions/audit_agent.py",
                "есть",
            ],
            [
                "Audit Tool Server :8100",
                "Система записи кейса: propose / select / download / ingest / ask / brief",
                "backend/ FastAPI — главный custom",
                "есть (норма)",
            ],
            [
                "Ollama (хост)",
                "Чат qwen3.6:35b, embeddings qwen3-embedding (не MiniLM)",
                "Нет",
                "есть",
            ],
            [
                "SearXNG :8080",
                "Добыча актов по allowlist РБ. В чате веб-поиск выключен",
                "searxng/settings.yml",
                "есть",
            ],
            [
                "Индекс кейса + sync Knowledge",
                "Чанки по статьям, ask с цитатами; опционально коллекция npa-{case}",
                "knowledge_flow.py, chunker.py",
                "есть",
            ],
            [
                "Саммари Word",
                "Карточки актов из фрагментов + обзор проверки, ссылки [n]",
                "brief_flow.py, brief_docx.py",
                "есть",
            ],
            [
                "React frontend/",
                "Лаборатория API. Не витрина продукта",
                "заморожен, profile lab",
                "не развиваем",
            ],
            [
                "DuckDB / тесты / WP / rules",
                "Факт, процедуры, черновики бумаг, детерминированные проверки",
                "план",
                "нет",
            ],
            [
                "Continue.dev / Pipelines-сервис",
                "IDE и отдельный orchestrator",
                "не нужны в v1",
                "не ставим",
            ],
        ],
    )

    add_heading_styled(doc, "1.2 Поток данных (как это работает сейчас)", 2)
    add_para(doc, "Аудитор живёт в одном окне http://localhost:3000, модель «Аудитор».")
    add_number(
        doc,
        "Описать проверку одной фразой: «Проверка аренды коммерческой недвижимости, аренда, валюта, НДС».",
    )
    add_number(
        doc,
        "Pipe вызывает POST /cases и POST /propose. Локальная LLM возвращает список актов с приоритетом 1/2/3 и обоснованием. Файлы ещё не качаются. В чат вшивается скрытая метка кейса.",
    )
    add_number(
        doc,
        "Аудитор: «утверждаю 1, 2, 4 плюс Инструкция НБРБ № 38» или «к 3 url https://pravo.by/…». Без явных id download запрещён.",
    )
    add_number(
        doc,
        "POST /select + POST /download. URL в порядке: ручная ссылка → curated known_sources → SearXNG (allowlist). Клиентский текст в поиск не уходит.",
    )
    add_number(
        doc,
        "HTML/PDF → очищенный текст → нарезка по «Статья N» → knowledge_index.json → (если есть ключ) sync в Open WebUI Knowledge.",
    )
    add_number(
        doc,
        "«саммари» — Word на 6–10 стр. из фрагментов, не из памяти модели. Вопросы по норме — RAG по библиотеке этого кейса.",
    )

    add_heading_styled(doc, "1.3 Папка кейса — система записи", 2)
    add_para(
        doc,
        "data/audit_cases/{case_id}/ : case.json, manifest.json, knowledge_raw/, knowledge_text/, "
        "summaries/, knowledge_index.json, library.zip. Позже: evidence/, tests/, wp/, trail/.",
        italic=True,
    )

    add_heading_styled(doc, "1.4 Принцип сборки", 2)
    add_para(
        doc,
        "Минимум своего кода. Максимум готовых инструментов. Пишем только то, чего нет на рынке: "
        "жизненный цикл кейса, домен РБ (allowlist, known_sources, нарезка НПА, норма ≠ факт), "
        "tools агента, audit trail. Не пишем чат, свою LLM, свой поисковик, второй фронт «потому что Open WebUI недостаточно красивый».",
    )
    add_para(
        doc,
        "Pipelines не содержат бизнес-логики: вызывают API сервера и показывают результат. "
        "Сервер не рисует UI. Native Tools (audit_case.py) — запасной путь: 35B ломает HITL.",
    )

    # ========== 2 ==========
    add_heading_styled(doc, "2. Архитектура (слои)", 1)
    add_para(
        doc,
        "Поверхность (готовое): Open WebUI — чат, Knowledge, citations, auth, Files.",
    )
    add_para(
        doc,
        "Клей агента: Pipe «Аудитор» — фазы библиотека → вопросы к НПА → (план) данные → тесты → WP. HITL в коде.",
    )
    add_para(
        doc,
        "Ядро (свой код): Audit Tool Server — кейс, HITL, добыча НПА, ingest, ask, brief; позже evidence, tests, WP, rules, trail.",
    )
    add_para(
        doc,
        "Исполнители: Ollama (LLM + embed), SearXNG (только добыча актов), диск кейса, Chroma внутри Open WebUI (индекс чата).",
    )

    add_heading_styled(doc, "2.1 Норма ≠ факт (архитектурный инвариант)", 2)
    add_table(
        doc,
        ["", "Норма", "Факт"],
        [
            ["Что", "НПА, инструкции НБРБ, внутренние политики", "Выгрузки, договоры, обороты"],
            ["Где", "knowledge_*, коллекция npa-{id}", "evidence/, DuckDB (план)"],
            ["Как отвечает", "RAG + цитата статьи", "tool / SQL / pandas / rule"],
            ["Риск ошибки", "галлюцинация статьи", "«похожий ряд в Excel»"],
        ],
    )
    add_para(
        doc,
        "Смешивать Excel клиента и ГК в одной Knowledge-коллекции нельзя. Вопросы «сколько начислено» "
        "через код, не через «похожий абзац».",
    )

    add_heading_styled(doc, "2.2 Allowlist поиска", 2)
    add_para(
        doc,
        "pravo.gov.by / pravo.by / etalonline.by, nbrb.by, minfin.gov.by, nalog.gov.by, "
        "government.by, president.gov.by. Defense in depth: фильтр и в SearXNG, и в сервере. "
        "Веб-поиск в чате выключен: иначе модель тащит чужую юрисдикцию и меню сайта в ответ.",
    )

    # ========== 3 ==========
    add_heading_styled(doc, "3. Как максимизировать качество ответов и работы", 1)
    add_para(
        doc,
        "Качество Cursor — это не «взять модель побольше». Это инженерия контекста: "
        "в окно модели попадает ровно тот файл и та статья, которые нужны, а факты "
        "считает код. Ниже — рычаги в порядке отдачи. Без пунктов A–D наращивать "
        "фазы проверки бессмысленно: аудитор не будет доверять тесту, если статья выдумана.",
    )

    add_heading_styled(doc, "3.1 Карта рычагов (что реально двигает качество)", 2)
    add_table(
        doc,
        ["Рычаг", "Аналог в Cursor", "Эффект", "Когда"],
        [
            [
                "A. Grounding: только утверждённая библиотека кейса",
                "Цитата из этого репо",
                "Критический. Убирает «вообще ГК» и РФ/IFRS",
                "сейчас уже; ужесточить промпт",
            ],
            [
                "B. Нарезка по статье + метаданные (акт, статья, URL, дата)",
                "Семантические чанки файла + путь",
                "Критический. «ст. 625» находится и цитируется",
                "ближайшие дни",
            ],
            [
                "C. Hybrid + rerank, не «похожий абзац»",
                "grep + embeddings",
                "Высокий. Номера статей, перефразы, меньше путаницы актов",
                "ближайшие дни",
            ],
            [
                "D. Eval: золотые вопросы после каждой смены RAG",
                "Тесты агента / eval harness",
                "Критический. Иначе крутят RAG «на глаз» и ломают доверие",
                "параллельно с B–C",
            ],
            [
                "E. Отказ как фича",
                "Модель не выдумывает API",
                "Высокий. «В библиотеке нет» важнее гладкого абзаца",
                "промпт + порог retrieval",
            ],
            [
                "F. Мнение из коробки (seed)",
                "Дефолтная модель, shortcuts",
                "Высокий. MiniLM / ctx=2048 убивают качество молча",
                "compose + засев Admin",
            ],
            [
                "G. HITL на список НПА (и позже на тесты)",
                "Apply / reject diff",
                "Высокий. Мусорный акт в индексе = мусорный ответ",
                "есть",
            ],
            [
                "H. Пакет типовой проверки",
                "Правила проекта + индекс",
                "Высокий. Повтор не с пустого промпта",
                "пилот «Аренда»",
            ],
            [
                "I. Факт через код, не RAG по Excel",
                "Терминал / линтер",
                "Критический для цифр",
                "после стабильной цитаты",
            ],
            [
                "J. Rules на курсы/лимиты",
                "Тесты в CI",
                "Средний, точечный",
                "1–2 правила на пилот",
            ],
            [
                "K. Trail: какой чанк ушёл в промпт",
                "История tool-вызовов",
                "Доверие + отладка",
                "сразу после eval",
            ],
            [
                "L. Модель 70B+ / reranker, когда железо тянет",
                "Лучшая модель в агенте",
                "Средний. Бесполезен, если retrieval врёт",
                "после зелёного eval",
            ],
        ],
    )

    add_heading_styled(doc, "3.2 A–E: качество цитаты (ядро доверия)", 2)
    add_para(
        doc,
        "Инвариант retrieval: модель видит только то, что попало в промпт. Если нужная статья "
        "не в top-k — модель её «не знает», даже если файл лежит в кейсе. Баги почти всегда "
        "не в LLM, а в extract / chunk / embed / top-k / num_ctx.",
        italic=True,
    )

    add_rich(
        doc,
        [
            ("1) Нормализация текста перед индексом. ", {"bold": True}),
            (
                "В Knowledge и в knowledge_index класть не сырой HTML с меню pravo.by, а markdown: "
                "заголовок акта, официальный URL, дата редакции, затем ## Статья N / ## Пункт N. "
                "Тогда сплиттер Open WebUI (Markdown Header Split) режет по той же границе, что и наш ARTICLE_RE. "
                "Не плодить вторую vector DB, пока этот путь не исчерпан.",
                {},
            ),
        ],
    )
    add_rich(
        doc,
        [
            ("2) Hybrid обязателен для «ст. 625». ", {"bold": True}),
            (
                "Семантика плохо ловит идентификаторы. BM25 ловит номер статьи как ключевое слово. "
                "Сейчас серверный /knowledge/ask — смесь keyword overlap + cosine, top_k=8; "
                "эмбеддятся не все чанки, а «релевантный» поднабор. Это слабее, чем hybrid Open WebUI "
                "(vector + BM25 + RRF + rerank). Цель: один контур качества — либо довести серверный ask "
                "до hybrid+rerank и метаданных статьи/URL в блоке «Откуда», либо сделать основным канал "
                "Open WebUI Knowledge после markdown-ingest, а серверный ask оставить контрольным.",
                {},
            ),
        ],
    )
    add_rich(
        doc,
        [
            ("3) Жёсткий отказ. ", {"bold": True}),
            (
                "В RAG-шаблоне Open WebUI уже правильно: «только фрагменты, не выдумывай, не РФ/IFRS». "
                "В ASK_SYSTEM сервера до сих пор лазейка: «можно добавить оговорку из своих знаний». "
                "Её нужно убрать. Нет фрагмента — «в утверждённой библиотеке этого нет», перечислить, "
                "где искали (имена актов). Калибровать промптом и порогом rerank, не «уверенностью %».",
                {},
            ),
        ],
    )
    add_rich(
        doc,
        [
            ("4) Цитата = файл в кейсе. ", {"bold": True}),
            (
                "В ответе: полное название акта, статья/пункт, дословная формулировка, URL. "
                "Клик/ссылка открывает тот же .txt/.html, что в knowledge_raw. Путаница «ГК vs инструкция НБРБ» "
                "лечится метаданными в чанке и шаблоном «называй документ полностью».",
                {},
            ),
        ],
    )
    add_rich(
        doc,
        [
            ("5) Золотой набор вопросов. ", {"bold": True}),
            (
                "20–30 пар вопрос–ожидание на пилотную тему (аренда). Четыре класса: (а) точная формулировка, "
                "которая есть в тексте; (б) статьи, которой нет — должен отказ; (в) запрос «ст. N» — BM25; "
                "(г) перефраз без номера — семантика. Скрипт eval_rag.py гонять после смены embedding, chunk, "
                "промпта. Без этого любая «улучшалка» — лотерея.",
                {},
            ),
        ],
    )
    add_rich(
        doc,
        [
            ("6) Окно контекста. ", {"bold": True}),
            (
                "chunk ~1500–2000, overlap 150–200, top-k 12–20, num_ctx модели 32k+, RAG_SYSTEM_CONTEXT=true. "
                "На 2048 токенах retrieved чанки обрежутся — будет казаться, что «embedding плохой». "
                "Full context глобально выключен: ГК в каждый ход не влезает. Точечно — на короткую инструкцию.",
                {},
            ),
        ],
    )

    add_heading_styled(doc, "3.3 F–H: ощущение продукта (коробка, не стенд)", 2)
    add_bullet(
        doc,
        "Одна кастомная модель «Аудитор»: промпт, Pipe, RAG, num_ctx уже стоят. Аудитор не выбирает между 15 моделями Ollama.",
        bold_prefix="Один персонаж. ",
    )
    add_bullet(
        doc,
        "compose up → один URL. Засев Pipe без ручной вставки в Admin — дыра коробки (сейчас Admin руками). "
        "На чистом volume hybrid/embedding уже в env; на живом volume — чеклист seed/openwebui.",
        bold_prefix="Засев. ",
    )
    add_bullet(
        doc,
        "Пилот «Аренда коммерческой недвижимости»: готовые keywords, known_sources (официальные URL ГК/НК/валютного закона уже есть), "
        "золотые вопросы, позже шаблон WP. Это контент, не код — сильный рычаг «коробочности», как правила репозитория в Cursor.",
        bold_prefix="Пакет проверки. ",
    )
    add_bullet(
        doc,
        "Ручная ссылка — первый класс. SearXNG поисковики банят (403). Порядок URL уже верный: manual → known_sources → SearXNG. "
        "Ошибки по-человечески: «вставьте URL с pravo.by», а не stack trace.",
        bold_prefix="Добыча актов. ",
    )
    add_bullet(
        doc,
        "Саммари в начале кейса — карточки «зачем этот акт» на 2 экрана, не 200 страниц кодекса. Уже есть команда «саммари». "
        "Качество карточки = качество отбора фрагментов по keywords проверки (не пересказ всего ГК).",
        bold_prefix="Не читай кодекс. ",
    )
    add_bullet(
        doc,
        "Фазы в коде Pipe, не свободный ReAct. 35B иначе вызовет download без «утверждаю». Качать/не качать решает аудитор. "
        "Когда цитата стабильна, 35B не нужна для выбора фазы — модель пишет текст по фрагментам.",
        bold_prefix="Агент, а не скрипт. ",
    )

    add_heading_styled(doc, "3.4 I–K: когда качество упирается не в RAG", 2)
    add_para(
        doc,
        "Cursor не эмбеддит вывод компилятора «на всякий случай» — он запускает компилятор. "
        "У нас то же разделение.",
    )
    add_bullet(
        doc,
        "Каталог evidence/ в кейсе. Загрузка файлом в кейс через tool или drop в папку. Не в коллекцию npa-*. "
        "Tool evidence_sql: DuckDB read_xlsx/csv, лимит строк, в trail — запрос и имя файла. Этот tool не ходит в SearXNG.",
        bold_prefix="Факт. ",
    )
    add_bullet(
        doc,
        "Модель данных теста: цель, критерий, источник нормы (id акта + статья), выборка, шаги, ожидаемый ok/finding. "
        "propose_tests → HITL select_tests. Шаблон WP банка (.docx): норма / факт / вывод / ссылка на файл. Статус всегда draft.",
        bold_prefix="Черновик, не подпись. ",
    )
    add_bullet(
        doc,
        "Простые YAML/код: курс НБРБ на дату, срок репатриации, лимит. Модель не пересчитывает «примерно». "
        "LLM только объясняет fail человеческим языком. 1–2 правила на пилот, не фреймворк на 40.",
        bold_prefix="Rules. ",
    )
    add_bullet(
        doc,
        "В case.json / trail/*.jsonl: какой file_id ушёл в Open WebUI, какой чанк и score попал в промпт, "
        "какой tool с какими аргументами, кто утвердил список. Строки выборки клиента — маскировать или не хранить.",
        bold_prefix="Trail. ",
    )
    add_bullet(
        doc,
        "QA-валидатор (как в плане 1.0, но на правильном месте): блокировать draft finding без evidence и без norm_reference. "
        "Это после тестов/WP, не вместо цитаты.",
        bold_prefix="Стоп-кран находок. ",
    )

    add_heading_styled(doc, "3.5 Модель и железо — после retrieval, не вместо", 2)
    add_para(
        doc,
        "Сейчас: qwen3.6:35b + qwen3-embedding. Это рабочий контур. 72B имеет смысл, когда золотые тесты "
        "уже зелёные и узкое место — рассуждение (программа проверки, формулировка finding), а не «не та статья». "
        "Reranker (bge-reranker-v2-m3) даёт больше, чем смена чат-модели, если GPU позволяет. "
        "Двухмодельный контур (маленькая классифицирует «норма/факт/процедура», большая отвечает) — когда одна "
        "сильная модель стабильно не путает фазу; сейчас фазу держит Pipe.",
    )
    add_para(
        doc,
        "Не делать: fine-tune на НПА (RAG + чанкер дешевле и честнее); мультиагенты «исследователь / скептик / писатель» "
        "(сначала один агент с HITL); native function calling как продуктовый путь на 35B.",
    )

    add_heading_styled(doc, "3.6 Что чинить в коде прямо сейчас (разрыв с планкой)", 2)
    add_para(
        doc,
        "Это не новый scope — это долг качества уже собранного контура нормы:",
    )
    add_table(
        doc,
        ["Разрыв", "Почему бьёт по качеству", "Что сделать"],
        [
            [
                "ASK_SYSTEM разрешает «знания модели»",
                "Противоречит RAG-шаблону и инварианту цитаты",
                "Убрать оговорку; единый шаблон отказа",
            ],
            [
                "В Knowledge уходит .txt без ## Статья N",
                "Сплиттер Open WebUI режет посередине нормы",
                "Нормализатор перед sync",
            ],
            [
                "ask: top_k=8, эмбед не всех чанков, нет статьи/URL в sources",
                "Пропуск нужного пункта, слабая сверка",
                "Hybrid, top-k 12–20, метаданные в «Откуда»",
            ],
            [
                "Нет eval_rag.py",
                "Регрессии невидимы",
                "20–30 золотых вопросов, гонять после смены RAG",
            ],
            [
                "Засев Pipe руками в Admin",
                "Коробка не «просто работает»",
                "Автозасев при compose up",
            ],
            [
                "known_sources тонкий",
                "SearXNG 403 → пустая библиотека",
                "Добить URL пилота «Аренда»",
            ],
            [
                "Два канала RAG (сервер и Open WebUI)",
                "Разное качество в зависимости от пути",
                "Один целевой канал + контрольный",
            ],
        ],
    )

    # ========== 4 ==========
    add_heading_styled(doc, "4. План наращивания (от факта к пилоту)", 1)
    add_para(
        doc,
        "Правило отбора работ. Прежде чем писать своё: (1) это уже есть в Open WebUI / Ollama / SearXNG / DuckDB / Word-шаблоне? "
        "(2) хватит ли tool-скрипта на 50–150 строк? (3) это домен проверки (кейс, HITL, НПА РБ, trail)? "
        "Нет → не делаем. Особенно экраны, админки, второй RAG.",
    )

    add_heading_styled(doc, "4.1 Этап 0–1. Коробка и мнение продукта", 2)
    add_para(
        doc,
        "Цель: compose up → открыл URL → чат с уже настроенным RAG. VISION/ARCHITECTURE/PLAN в репо — источник истины. "
        "Frontend в profile lab. Один вход = Open WebUI.",
    )
    add_bullet(doc, "Сделано: compose (Open WebUI, SearXNG, backend), seed RAG/hybrid в env, документ аудитора (AUDITOR.md), Pipe с HITL.")
    add_bullet(doc, "Осталось: корневой .env.example; страница «поставил → работает» (для банка — офлайн-бандл образов/весов); автозасев Pipe.")

    add_heading_styled(doc, "4.2 Этап 2. Аудитор живёт в чате (гипотеза «фронт не нужен»)", 2)
    add_para(
        doc,
        "Сделано: поток библиотеки НПА из чата, жёсткий HITL, extra_titles, ручной URL, саммари, ask. "
        "Решение по React — после того, как список из 15 актов удобен или нет. Если галочки в тексте больные — "
        "не портал, а одна HTMX-страница только на select. Новые шаги 3–6 в React не класть.",
    )

    add_heading_styled(doc, "4.3 Этап 3. Качество НПА — главный рычаг ближайших недель", 2)
    add_para(
        doc,
        "Цель: цитата стабильно из нужной статьи, отказ если статьи нет. Это ядро доверия. "
        "Сюда входят пункты 3.2 и 3.6. Выход: можно показать службе внутреннего аудита ответ с цитатой и не краснеть. "
        "Не делать на этом этапе: Postgres, агент на 7 фаз, новый фронт.",
    )

    add_heading_styled(doc, "4.4 Этап 4. Кейс как система записи", 2)
    add_bullet(doc, "Список кейсов в чате: «открой кейс X», привязать Knowledge. Файловый store; SQLite — только если больно искать.")
    add_bullet(doc, "Статусы v1: created → proposed → selected → ready. Не BPM «согласовано начальником».")
    add_bullet(doc, "Режим airgap: SearXNG выключен, только manual URL и known_sources.")

    add_heading_styled(doc, "4.5 Этап 5. Данные клиента без RAG по Excel", 2)
    add_para(
        doc,
        "evidence/ + evidence_sql (DuckDB). 2–3 типовых запроса в system prompt (сумма, период, группировка). "
        "Выход: норма и факт отвечают разными механизмами.",
    )

    add_heading_styled(doc, "4.6 Этап 6. Черновики тестов и WP", 2)
    add_para(
        doc,
        "Не «генератор отчёта». Аудитор копирует черновик в свой контур рабочей документации. "
        "QA: finding без нормы и без evidence не собирается.",
    )

    add_heading_styled(doc, "4.7 Этап 7. Rules там, где LLM врать дороже всего", 2)
    add_para(doc, "run_rule → pass/fail + explanation. Пайплайн: сначала rule, потом текст.")

    add_heading_styled(doc, "4.8 Этап 8. Полировка коробки", 2)
    add_bullet(doc, "Инсталлятор: GPU/RAM, pull моделей, healthcheck, «откройте http://…».")
    add_bullet(doc, "Пакет пилота «Аренда». Выкинуть лабораторный фронт из коробки.")
    add_bullet(doc, "Пилот с 2 аудиторами на 1–2 реальных проверках. Feedback в промпты, known_sources и eval — не в новый UI.")

    add_heading_styled(doc, "4.9 Если времени мало — порядок ценности", 2)
    add_number(doc, "Коробка и засев RAG (этап 1).")
    add_number(doc, "Цитаты и eval (этап 3) — можно частично параллельно.")
    add_number(doc, "Удобство HITL в чате (этап 2), решение по фронту.")
    add_number(doc, "DuckDB по одной реальной выгрузке.")
    add_number(doc, "Один шаблон WP отдела.")
    add_number(doc, "Остальное.")
    add_para(
        doc,
        "Писать React-портал при неготовой коробке и сырой цитате — самый дорогой способ не получить продукт. "
        "План 1.0 (сразу 10 фаз + data_profile + Continue.dev) сознательно не возвращаем: это расползание до качества Cursor.",
        italic=True,
    )

    # ========== 5 ==========
    add_heading_styled(doc, "5. Целевой UX аудитора", 1)
    add_para(doc, "Сообщение:", italic=True)
    add_para(
        doc,
        "Проверка аренды коммерческой недвижимости, аренда, валюта, НДС",
        bold=True,
    )
    add_para(doc, "Цикл в одном чате:", italic=True)
    add_bullet(doc, "Список актов (модель предлагает) → аудитор утверждает номера и/или даёт URL.")
    add_bullet(doc, "Скачивание → «документы» / «саммари» → Word со ссылками [n].")
    add_bullet(doc, "Вопросы: «Можно ли устанавливать арендную плату в валюте для резидентов РБ?» → ответ + блок «Откуда».")
    add_bullet(doc, "Нет акта в библиотеке — отказ → доутвердить акт или вставить ссылку, не сочинять.")
    add_bullet(doc, "Позже: «профилируй ops.xlsx» → SQL → «предложи тесты» → HITL → «черновик WP по тесту 3».")
    add_para(
        doc,
        "Новая проверка — новый чат. Не смешивать две темы. Не прикреплять клиентский Excel «чтобы модель посмотрела» — смешает закон и цифру. "
        "Полные фразы и разбор сбоев: docs/AUDITOR.md.",
    )

    # ========== 6 ==========
    add_heading_styled(doc, "6. Конфиденциальность и контур банка", 1)
    add_bullet(doc, "Ollama и Open WebUI без облачных API-ключей. Запрещены OpenAI / Claude / Gemini / облачные embeddings / Tavily.")
    add_bullet(doc, "Клиентские данные только в /audit_cases/{case_id}/evidence — не в общий Knowledge.")
    add_bullet(doc, "В query SearXNG не подставлять счета, ФИО, суммы, формулировки проверки.")
    add_bullet(doc, "Tools к evidence/ не имеют права ходить в интернет.")
    add_bullet(doc, "Air-gap: образы и веса моделей файлами, не Hugging Face из банка.")
    add_bullet(doc, "RBAC auditor / reviewer / admin — не в v1, если один аудитор на машине. Четыре глаза — идея после пилота.")
    add_bullet(doc, "SQL к АБС (если появится) — только read-only replica по политике банка.")

    # ========== 7 ==========
    add_heading_styled(doc, "7. Анти-цели (пока «как Cursor» не достигнуто)", 1)
    add_table(
        doc,
        ["Не делать", "Почему"],
        [
            ["Свой чат / SPA «портал аудитора»", "Дубль Open WebUI, вечный второй сорт"],
            ["Форк Open WebUI «чуть UX»", "Навсегда отстанем от апстрима"],
            ["10 автономных фаз без стабильной цитаты", "Автопилот по выдуманным нормам"],
            ["RAG по Excel в куче с ГК", "Галлюцинации фактов"],
            ["Автоматическое юридическое заключение", "Ломает инвариант продукта"],
            ["Веб-поиск дефолтом чата", "Утечки и чужая юрисдикция"],
            ["Мультиагенты, fine-tune, Kubernetes, Kafka", "Сложность вместо качества ответа"],
            ["Continue.dev / Aider в коробке v1", "IDE для правки rules — после стабильного чата"],
            ["Смена чат-модели до зелёных золотых тестов", "Маскирует сломанный retrieval"],
        ],
    )

    # ========== 8 ==========
    add_heading_styled(doc, "8. Чеклисты", 1)
    add_heading_styled(doc, "8.1 Коробка стоит", 2)
    add_bullet(doc, "[x] docker compose: Open WebUI, SearXNG, Audit Tool Server")
    add_bullet(doc, "[x] Ollama на хосте: qwen3.6:35b + qwen3-embedding (не MiniLM)")
    add_bullet(doc, "[x] Веб-поиск в чате выключен; SearXNG только для download")
    add_bullet(doc, "[x] Нет облачных ключей")
    add_bullet(doc, "[ ] Pipe «Аудитор» засеян без ручного Admin (сейчас — руками, AGENT.md)")
    add_bullet(doc, "[ ] num_ctx модели 32k+ проверен на живом volume")

    add_heading_styled(doc, "8.2 Цитата не стыдная", 2)
    add_bullet(doc, "[x] HITL: download только после «утверждаю»")
    add_bullet(doc, "[x] RAG-шаблон: только фрагменты, не выдумывай, не РФ/IFRS")
    add_bullet(doc, "[ ] ASK_SYSTEM без лазейки «из своих знаний»")
    add_bullet(doc, "[ ] .txt → markdown ## Статья N перед sync")
    add_bullet(doc, "[ ] Четыре золотых теста на одной инструкции НБРБ проходят")
    add_bullet(doc, "[ ] 20–30 вопросов пакета «Аренда» в eval_rag.py")
    add_bullet(doc, "[ ] На вопросе вне библиотеки — отказ + список актов, где искали")

    add_heading_styled(doc, "8.3 Готово показать аудитору без репозитория", 2)
    add_bullet(doc, "[ ] Один URL, модель «Аудитор», без настройки RAG руками")
    add_bullet(doc, "[x] Список НПА нельзя скачать без утверждения")
    add_bullet(doc, "[x] «саммари» отдаёт Word")
    add_bullet(doc, "[ ] Вопрос по акту — с цитатой из файла кейса")
    add_bullet(doc, "[ ] Нет второго окна Vite, без которого нельзя работать")
    add_bullet(doc, "[ ] Повтор той же темы начинается с пакета, не с пустого промпта")

    add_heading_styled(doc, "8.4 Пилот (после этапов 5–6)", 2)
    add_bullet(doc, "[ ] evidence_sql на одной реальной выгрузке банка")
    add_bullet(doc, "[ ] Черновик WP по шаблону отдела, статус draft")
    add_bullet(doc, "[ ] QA блокирует finding без нормы и evidence")
    add_bullet(doc, "[ ] Trail восстановим")
    add_bullet(doc, "[ ] 1–2 проверки с живыми аудиторами")

    # ========== 9 ==========
    add_heading_styled(doc, "9. Критерии «мы близко к Cursor»", 1)
    add_para(doc, "Не «есть GPT в чате», а:")
    add_bullet(doc, "Человек без этого репо поднимает систему одной командой и задаёт вопрос по акту с цитатой.")
    add_bullet(doc, "На вопросе, которого нет в библиотеке, модель отказывается, а не сочиняет статью.")
    add_bullet(doc, "Список НПА нельзя скачать без утверждения аудитора.")
    add_bullet(doc, "Excel клиента не всплывает в ответе про закон; закон не всплывает в ответе про сумму без tool.")
    add_bullet(doc, "Нет второго окна «настоящего UI», без которого нельзя работать.")
    add_bullet(doc, "Повторная проверка той же темы начинается с пакета, а не с пустого промпта.")
    add_bullet(doc, "Аудитор тратит время на суждение и выборку, а не на поиск редакции на pravo.by и перечитывание кодекса.")

    # ========== 10 ==========
    add_heading_styled(doc, "10. Источники истины в репозитории", 1)
    add_para(
        doc,
        "Этот Word — продуктовый план для команды и пилота. Детали кода живут рядом, чтобы не разъехаться:",
    )
    add_bullet(doc, "docs/VISION.md — что продаём.")
    add_bullet(doc, "docs/ARCHITECTURE.md — сборка, не форк.")
    add_bullet(doc, "docs/PLAN.md — этапы и чеклисты разработки.")
    add_bullet(doc, "docs/AUDITOR.md — как аудитор работает в чате.")
    add_bullet(doc, "docs/RAG_для_разработчика.md — retrieval, грабли, золотые тесты.")
    add_bullet(doc, "docs/START.md — что писать на этой неделе.")
    add_bullet(doc, "docs/IDEAS.md — бэклог, не обязательства.")
    add_bullet(doc, "seed/openwebui/ — промпт, RAG-шаблон, Pipe.")
    add_para(
        doc,
        "Сопутствующий PDF в той же папке (гипотезы / 10 фаз) остаётся справочником по будущим фазам проверки. "
        "Фазы данных и отчёта включаются только после зелёного этапа 3. Алгоритм «агент сам проходит всё без человека» "
        "не является целью v1: Cursor тоже не коммитит в main без вас.",
        italic=True,
    )

    # ========== 11 ==========
    add_heading_styled(doc, "11. Итог", 1)
    add_para(
        doc,
        "Путь: готовый стек уже стоит → довести цитату до доверия (markdown-статьи, hybrid, отказ, eval, пакет «Аренда») "
        "→ коробка без ручного Admin → факт через DuckDB → черновики тестов/WP и 1–2 rules → пилот. "
        "Качество, близкое к Cursor, даёт не магия 72B и не пятнадцать tools в первую неделю, "
        "а цикл «утверждённый контекст → tool/retrieval → проверка → отказ, если доказательства нет».",
    )
    add_callout(
        doc,
        "Ближайший спринт качества: убрать лазейку «из знаний модели», нормализовать статьи, "
        "посадить золотые вопросы, добить known_sources пилота. Пока эти четыре вещи красные — "
        "не начинать портал, Pipelines-оркестратор и автогенерацию findings.",
    )

    footer = doc.add_paragraph()
    p_style(footer, space_before=16, space_after=0)
    r = footer.add_run(
        "audit-tools  ·  версия документа 2.0  ·  24.08.2026  ·  "
        "предыдущая 1.0 (20.08.2026) сохранена по смыслу в git / этом файле как заменённая"
    )
    set_run_font(r, size=8, italic=True, color=MUTED)

    target = OUT
    try:
        doc.save(OUT)
    except PermissionError:
        target = OUT_FALLBACK
        doc.save(OUT_FALLBACK)
    print("wrote", target, "bytes", target.stat().st_size)


if __name__ == "__main__":
    build()
